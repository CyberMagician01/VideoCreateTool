import json
from io import BytesIO
import os
import re
import sqlite3
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional
import warnings

import requests
from dotenv import load_dotenv
from flask import Flask, jsonify, render_template, request, send_file
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import simpleSplit
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas


warnings.filterwarnings("ignore", message="Unverified HTTPS request")
load_dotenv()


QINIU_AK = os.getenv("QINIU_AK", os.getenv("AK", "")).strip()
QINIU_SK = os.getenv("QINIU_SK", os.getenv("SK", "")).strip()
QINIU_TEXT_API_KEY = os.getenv(
    "QINIU_TEXT_API_KEY",
    os.getenv("text_api_key", os.getenv("OPENAI_API_KEY", "")),
).strip()
QINIU_VIDEO_API_KEY = os.getenv("QINIU_VIDEO_API_KEY", os.getenv("video_api_key", "")).strip()
QINIU_LLM_MODEL = os.getenv("QINIU_LLM_MODEL", os.getenv("QWEN_MODEL", "qwen-plus"))
QINIU_LLM_BASE_URL = os.getenv(
    "QINIU_LLM_BASE_URL",
    os.getenv("QWEN_BASE_URL", os.getenv("OPENAI_BASE_URL", "")),
).strip()
QINIU_VIDEO_MODEL = os.getenv("QINIU_VIDEO_MODEL", os.getenv("WAN_MODEL", "wan2.6-t2v"))
QINIU_VIDEO_BASE_URL = os.getenv("QINIU_VIDEO_BASE_URL", os.getenv("WAN_BASE_URL", "")).strip()
QINIU_VIDEO_CREATE_PATH = os.getenv(
    "QINIU_VIDEO_CREATE_PATH", "/services/aigc/video-generation/video-synthesis"
).strip()
QINIU_VIDEO_TASK_PATH_TEMPLATE = os.getenv("QINIU_VIDEO_TASK_PATH_TEMPLATE", "/tasks/{task_id}").strip()
QINIU_VIDU_Q3_TEXT_TO_VIDEO_PATH = os.getenv(
    "QINIU_VIDU_Q3_TEXT_TO_VIDEO_PATH", "/queue/fal-ai/vidu/q3/text-to-video/turbo"
).strip()
QINIU_VIDU_Q3_IMAGE_TO_VIDEO_PATH = os.getenv(
    "QINIU_VIDU_Q3_IMAGE_TO_VIDEO_PATH", "/queue/fal-ai/vidu/q3/image-to-video/turbo"
).strip()
QINIU_VIDU_Q3_START_END_TO_VIDEO_PATH = os.getenv(
    "QINIU_VIDU_Q3_START_END_TO_VIDEO_PATH", "/queue/fal-ai/vidu/q3/start-end-to-video/turbo"
).strip()
QINIU_WEB_SEARCH_PATH = os.getenv("QINIU_WEB_SEARCH_PATH", "/search/web").strip()
QINIU_LLM_FALLBACK_MODELS = [
    item.strip()
    for item in os.getenv("QINIU_LLM_FALLBACK_MODELS", "").split(",")
    if item.strip()
]
QINIU_ENABLE_ASYNC_HEADER = os.getenv("QINIU_ENABLE_ASYNC_HEADER", "true").strip().lower() in {
    "1",
    "true",
    "yes",
    "on",
}

DEFAULT_PROVIDER = os.getenv("DEFAULT_PROVIDER", "qiniu")


MODEL_PROVIDERS = {
    "qiniu": {
        "name": "七牛云大模型网关",
        "model": QINIU_LLM_MODEL,
        "base_url": QINIU_LLM_BASE_URL,
        "type": "qiniu_openai_compatible",
        "requires_qiniu_credential": True,
    }
}


app = Flask(__name__, template_folder="templates", static_folder="static")


DATA_DIR = Path(__file__).resolve().parent / "data"
DB_PATH = DATA_DIR / "projects.db"


def _utc_now_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _default_project_state() -> Dict[str, Any]:
    return {
        "story_card": None,
        "workshop": None,
        "storyboard": None,
        "video_lab": {
            "script": "",
            "prompt": "",
            "task_id": "",
            "task_status": "",
            "video_url": "",
            "auto_poll": True,
            "last_check_time": "",
            "long_segments": [],
            "total_duration": 0,
            "filename_prefix": "",
        },
    }


def _as_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def _string_list(value: Any) -> List[str]:
    if isinstance(value, list):
        result: List[str] = []
        for item in value:
            text = _as_text(item)
            if text:
                result.append(text)
        return result
    text = _as_text(value)
    return [text] if text else []


def _safe_int(value: Any, default: int = 0, *, minimum: Optional[int] = None) -> int:
    try:
        parsed = int(value)
    except (TypeError, ValueError):
        parsed = default
    if minimum is not None:
        parsed = max(minimum, parsed)
    return parsed


def _first_present(data: Dict[str, Any], *keys: str) -> Any:
    for key in keys:
        if key in data:
            value = data.get(key)
            if value not in (None, "", [], {}):
                return value
    return None


def _derive_storyboard_prompt(shot: Dict[str, Any]) -> str:
    parts = [
        _as_text(shot.get("shot_type")),
        _as_text(shot.get("camera_movement")),
        _as_text(shot.get("visual_description")),
        _as_text(shot.get("dialogue_or_sfx")),
    ]
    return " | ".join(part for part in parts if part)


def _normalize_story_card(story_card: Any) -> Optional[Dict[str, Any]]:
    if not isinstance(story_card, dict):
        return None

    normalized = {
        "logline": _as_text(story_card.get("logline")),
        "theme": _as_text(story_card.get("theme")),
        "tone": _as_text(story_card.get("tone")),
        "structure_template": _as_text(story_card.get("structure_template")),
        "core_conflict": _as_text(story_card.get("core_conflict")),
        "anchor_points": _string_list(story_card.get("anchor_points")),
        "hook": _as_text(story_card.get("hook")),
        "ending_type": _as_text(story_card.get("ending_type")),
    }
    if any(
        [
            normalized["logline"],
            normalized["theme"],
            normalized["tone"],
            normalized["structure_template"],
            normalized["core_conflict"],
            normalized["anchor_points"],
            normalized["hook"],
            normalized["ending_type"],
        ]
    ):
        return normalized
    return None


def _normalize_story_engine_result(result: Any) -> Dict[str, Any]:
    if not isinstance(result, dict):
        return {"story_card": None, "next_questions": []}
    return {
        "story_card": _normalize_story_card(result.get("story_card") if "story_card" in result else result),
        "next_questions": _string_list(result.get("next_questions")),
    }


def _normalize_character(character: Any) -> Optional[Dict[str, Any]]:
    if not isinstance(character, dict):
        return None
    normalized = {
        "name": _as_text(_first_present(character, "name", "character_name")),
        "tags": _string_list(_first_present(character, "tags", "labels")),
        "motivation": _as_text(_first_present(character, "motivation", "goal")),
        "arc": _as_text(_first_present(character, "arc", "character_arc")),
    }
    if normalized["name"]:
        return normalized
    return None


def _normalize_relationship(relationship: Any) -> Optional[Dict[str, Any]]:
    if not isinstance(relationship, dict):
        return None
    normalized = {
        "from": _as_text(_first_present(relationship, "from", "source", "from_character")),
        "to": _as_text(_first_present(relationship, "to", "target", "to_character")),
        "type": _as_text(_first_present(relationship, "type", "relationship", "relation")),
        "tension": _as_text(_first_present(relationship, "tension", "conflict")),
    }
    if normalized["from"] and normalized["to"]:
        return normalized
    return None


def _normalize_plot_node(node: Any, index: int) -> Optional[Dict[str, Any]]:
    if not isinstance(node, dict):
        return None
    normalized = {
        "id": _as_text(_first_present(node, "id", "node_id")) or f"N{index}",
        "template_stage": _as_text(_first_present(node, "template_stage", "phase", "stage")),
        "summary": _as_text(_first_present(node, "summary", "plot", "content", "scene_summary")),
        "location": _as_text(_first_present(node, "location", "scene_location")),
        "action_draft": _as_text(_first_present(node, "action_draft", "action", "action_description")),
        "dialogue_draft": _string_list(_first_present(node, "dialogue_draft", "dialogue", "dialogues")),
        "emotion_shift": _as_text(_first_present(node, "emotion_shift", "emotional_shift", "emotion")),
        "consistency_check": _as_text(_first_present(node, "consistency_check", "logic_check", "consistency")),
    }
    if any(
        [
            normalized["template_stage"],
            normalized["summary"],
            normalized["location"],
            normalized["action_draft"],
            normalized["dialogue_draft"],
            normalized["emotion_shift"],
            normalized["consistency_check"],
        ]
    ):
        return normalized
    return None


def _normalize_card_wall_group(group: Any) -> Optional[Dict[str, Any]]:
    if not isinstance(group, dict):
        return None
    normalized = {
        "group": _as_text(_first_present(group, "group", "name", "title")),
        "node_ids": _string_list(_first_present(group, "node_ids", "ids")),
    }
    if normalized["group"] or normalized["node_ids"]:
        return normalized
    return None


def _normalize_workshop_result(workshop: Any) -> Optional[Dict[str, Any]]:
    if not isinstance(workshop, dict):
        return None

    characters = [item for item in (_normalize_character(c) for c in workshop.get("characters", [])) if item]
    relationships = [
        item for item in (_normalize_relationship(r) for r in workshop.get("relationships", [])) if item
    ]
    plot_nodes = [
        item for item in (_normalize_plot_node(node, idx + 1) for idx, node in enumerate(workshop.get("plot_nodes", [])))
        if item
    ]
    available_node_ids = {node["id"] for node in plot_nodes}
    timeline_view = [
        node_id for node_id in _string_list(workshop.get("timeline_view")) if node_id in available_node_ids
    ]
    if not timeline_view:
        timeline_view = [node["id"] for node in plot_nodes]
    card_wall_groups = [
        item for item in (_normalize_card_wall_group(group) for group in workshop.get("card_wall_groups", [])) if item
    ]

    if any([characters, relationships, plot_nodes, timeline_view, card_wall_groups]):
        return {
            "characters": characters,
            "relationships": relationships,
            "plot_nodes": plot_nodes,
            "timeline_view": timeline_view,
            "card_wall_groups": card_wall_groups,
        }
    return None


def _normalize_storyboard_shot(shot: Any, index: int) -> Optional[Dict[str, Any]]:
    if not isinstance(shot, dict):
        return None
    normalized = {
        "shot_id": _as_text(_first_present(shot, "shot_id", "id")) or f"S{index}",
        "related_node_id": _as_text(_first_present(shot, "related_node_id", "node_id", "plot_node_id")),
        "shot_type": _as_text(_first_present(shot, "shot_type", "camera_size")),
        "camera_movement": _as_text(_first_present(shot, "camera_movement", "movement", "camera_motion")),
        "visual_description": _as_text(
            _first_present(shot, "visual_description", "visual", "image_description", "description")
        ),
        "dialogue_or_sfx": _as_text(_first_present(shot, "dialogue_or_sfx", "dialogue", "sound_design", "audio")),
        "duration_sec": _safe_int(_first_present(shot, "duration_sec", "duration", "estimated_duration"), 4, minimum=1),
        "shooting_note": _as_text(_first_present(shot, "shooting_note", "note", "production_note")),
        "prompt_draft": _as_text(_first_present(shot, "prompt_draft", "video_prompt", "prompt", "visual_prompt")),
    }
    if not normalized["prompt_draft"]:
        normalized["prompt_draft"] = _derive_storyboard_prompt(normalized)
    if any(
        [
            normalized["related_node_id"],
            normalized["shot_type"],
            normalized["camera_movement"],
            normalized["visual_description"],
            normalized["dialogue_or_sfx"],
            normalized["prompt_draft"],
        ]
    ):
        return normalized
    return None


def _normalize_storyboard_result(storyboard: Any) -> Optional[Dict[str, Any]]:
    if not isinstance(storyboard, dict):
        return None

    storyboards = [
        item
        for item in (
            _normalize_storyboard_shot(shot, idx + 1)
            for idx, shot in enumerate(storyboard.get("storyboards", []))
        )
        if item
    ]
    estimated_total_duration_sec = _safe_int(
        storyboard.get("estimated_total_duration_sec"),
        sum(item["duration_sec"] for item in storyboards),
        minimum=0,
    )
    export_ready_checklist = _string_list(storyboard.get("export_ready_checklist"))

    if storyboards or export_ready_checklist or estimated_total_duration_sec:
        return {
            "storyboards": storyboards,
            "estimated_total_duration_sec": estimated_total_duration_sec,
            "export_ready_checklist": export_ready_checklist,
        }
    return None


def _normalize_video_segment(segment: Any, index: int) -> Optional[Dict[str, Any]]:
    if not isinstance(segment, dict):
        return None
    normalized = {
        "index": _safe_int(segment.get("index"), index, minimum=1),
        "duration": _safe_int(segment.get("duration"), 0, minimum=0),
        "prompt": _as_text(segment.get("prompt")),
        "task_id": _as_text(segment.get("task_id")),
        "task_status": _as_text(segment.get("task_status")),
        "video_url": _as_text(_first_present(segment, "video_url", "url")),
    }
    if any(normalized.values()):
        return normalized
    return None


def _normalize_video_lab_state(video_lab: Any) -> Dict[str, Any]:
    base = _default_project_state()["video_lab"]
    if not isinstance(video_lab, dict):
        return base

    segments = [
        item
        for item in (
            _normalize_video_segment(segment, idx + 1)
            for idx, segment in enumerate(video_lab.get("long_segments", []))
        )
        if item
    ]
    return {
        "script": _as_text(video_lab.get("script")),
        "prompt": _as_text(video_lab.get("prompt")),
        "task_id": _as_text(video_lab.get("task_id")),
        "task_status": _as_text(video_lab.get("task_status")),
        "video_url": _as_text(_first_present(video_lab, "video_url", "url")),
        "auto_poll": bool(video_lab.get("auto_poll", True)),
        "last_check_time": _as_text(video_lab.get("last_check_time")),
        "long_segments": segments,
        "total_duration": _safe_int(video_lab.get("total_duration"), 0, minimum=0),
        "filename_prefix": _as_text(video_lab.get("filename_prefix")),
    }


def _normalize_project_state(state: Any) -> Dict[str, Any]:
    return {
        "story_card": _normalize_story_card(state.get("story_card")) if isinstance(state, dict) else None,
        "workshop": _normalize_workshop_result(state.get("workshop")) if isinstance(state, dict) else None,
        "storyboard": _normalize_storyboard_result(state.get("storyboard")) if isinstance(state, dict) else None,
        "video_lab": _normalize_video_lab_state(state.get("video_lab") if isinstance(state, dict) else None),
    }


def _normalize_command_result(result: Any) -> Dict[str, Any]:
    if not isinstance(result, dict):
        return {
            "command_understanding": "",
            "updated_state": {},
            "consistency_report": [],
            "suggestions": [],
        }

    updated_state = result.get("updated_state", {})
    normalized_updated_state: Dict[str, Any] = {}
    if isinstance(updated_state, dict):
        story_card = _normalize_story_card(updated_state.get("story_card"))
        workshop = _normalize_workshop_result(updated_state.get("workshop"))
        storyboard = _normalize_storyboard_result(updated_state.get("storyboard"))
        if story_card is not None:
            normalized_updated_state["story_card"] = story_card
        if workshop is not None:
            normalized_updated_state["workshop"] = workshop
        if storyboard is not None:
            normalized_updated_state["storyboard"] = storyboard

    return {
        "command_understanding": _as_text(result.get("command_understanding")),
        "updated_state": normalized_updated_state,
        "consistency_report": _string_list(result.get("consistency_report")),
        "suggestions": _string_list(result.get("suggestions")),
    }


def _get_db_conn() -> sqlite3.Connection:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def _init_projects_db() -> None:
    with _get_db_conn() as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS projects (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                creator TEXT DEFAULT '',
                description TEXT DEFAULT '',
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL,
                cover_image TEXT DEFAULT '',
                last_provider TEXT DEFAULT '',
                deleted INTEGER DEFAULT 0
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS project_states (
                project_id INTEGER PRIMARY KEY,
                state_json TEXT NOT NULL,
                updated_at TEXT NOT NULL,
                FOREIGN KEY(project_id) REFERENCES projects(id)
            )
            """
        )
        conn.commit()


def _row_to_project_meta(row: sqlite3.Row) -> Dict[str, Any]:
    return {
        "id": row["id"],
        "name": row["name"],
        "creator": row["creator"] or "",
        "description": row["description"] or "",
        "created_at": row["created_at"],
        "updated_at": row["updated_at"],
        "cover_image": row["cover_image"] or "",
        "last_provider": row["last_provider"] or "",
    }


def _create_project(
    conn: sqlite3.Connection,
    *,
    name: str,
    creator: str = "",
    description: str = "",
    state: Optional[Dict[str, Any]] = None,
    cover_image: str = "",
    last_provider: str = "",
) -> Dict[str, Any]:
    now = _utc_now_iso()
    state_payload = _normalize_project_state(state)
    cur = conn.execute(
        """
        INSERT INTO projects(name, creator, description, created_at, updated_at, cover_image, last_provider, deleted)
        VALUES (?, ?, ?, ?, ?, ?, ?, 0)
        """,
        (name.strip(), creator.strip(), description.strip(), now, now, cover_image, last_provider),
    )
    project_id = cur.lastrowid
    conn.execute(
        """
        INSERT INTO project_states(project_id, state_json, updated_at)
        VALUES (?, ?, ?)
        """,
        (project_id, json.dumps(state_payload, ensure_ascii=False), now),
    )
    conn.commit()
    row = conn.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
    return _row_to_project_meta(row)


def _list_projects(conn: sqlite3.Connection) -> List[Dict[str, Any]]:
    rows = conn.execute(
        """
        SELECT * FROM projects
        WHERE deleted = 0
        ORDER BY datetime(updated_at) DESC, id DESC
        """
    ).fetchall()
    return [_row_to_project_meta(r) for r in rows]


def _ensure_default_project(conn: sqlite3.Connection) -> Dict[str, Any]:
    projects = _list_projects(conn)
    if projects:
        return projects[0]
    return _create_project(conn, name="未命名项目")


def _get_project_with_state(conn: sqlite3.Connection, project_id: int) -> Optional[Dict[str, Any]]:
    row = conn.execute(
        "SELECT * FROM projects WHERE id = ? AND deleted = 0",
        (project_id,),
    ).fetchone()
    if not row:
        return None

    state_row = conn.execute(
        "SELECT state_json FROM project_states WHERE project_id = ?",
        (project_id,),
    ).fetchone()

    state_obj: Dict[str, Any] = _default_project_state()
    if state_row and state_row["state_json"]:
        try:
            parsed = json.loads(state_row["state_json"])
            if isinstance(parsed, dict):
                state_obj = _normalize_project_state(parsed)
        except json.JSONDecodeError:
            state_obj = _default_project_state()

    return {
        "project": _row_to_project_meta(row),
        "state": state_obj,
    }


_init_projects_db()


@dataclass
class AgentRequest:
    stage: str
    payload: Dict[str, Any]
    provider: str = DEFAULT_PROVIDER


def _extract_json(text: str) -> Dict[str, Any]:
    """Extract the first JSON object from model output and parse it safely."""
    text = text.strip()

    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?", "", text).strip()
        text = re.sub(r"```$", "", text).strip()

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    match = re.search(r"\{[\s\S]*\}", text)
    if not match:
        raise ValueError("Model did not return JSON.")

    return json.loads(match.group(0))


def _resolve_url(base_url: str, path_or_url: str) -> str:
    text = str(path_or_url or "").strip()
    if text.startswith("http://") or text.startswith("https://"):
        return text

    base = str(base_url or "").strip().rstrip("/")
    if not base:
        raise RuntimeError("QINIU base url is missing in .env")

    if not text:
        return base

    if not text.startswith("/"):
        text = f"/{text}"
    return f"{base}{text}"


def _build_url_candidates(base_url: str, path_or_url: str) -> List[str]:
    """Build URL candidates for gateways that differ on /v1 vs /api/v1 prefix."""
    primary = _resolve_url(base_url, path_or_url)
    candidates = [primary]

    base = str(base_url or "").strip().rstrip("/")
    path = str(path_or_url or "").strip()

    if path and not path.startswith("/"):
        path = f"/{path}"

    if base.endswith("/api/v1"):
        alt_base = base[: -len("/api/v1")] + "/v1"
        candidates.append(f"{alt_base}{path}")
    elif base.endswith("/v1"):
        alt_base = base[: -len("/v1")] + "/api/v1"
        candidates.append(f"{alt_base}{path}")

    deduped: List[str] = []
    for url in candidates:
        if url not in deduped:
            deduped.append(url)
    return deduped


def _is_not_found_or_method_not_allowed(err: requests.HTTPError) -> bool:
    response = err.response
    if response is None:
        return False

    if response.status_code in {404, 405}:
        return True

    detail = response.text.lower()
    return ("not found" in detail) or ("method not allowed" in detail)


def _build_qiniu_aksk_headers(*, include_content_type: bool = True) -> Dict[str, str]:
    if not QINIU_AK or not QINIU_SK:
        raise RuntimeError("QINIU_AK or QINIU_SK is missing in .env")

    headers = {
        "X-Qiniu-AK": QINIU_AK,
        "X-Qiniu-SK": QINIU_SK,
        # 部分七牛网关使用自定义授权头，这里统一附加，具体格式可按网关文档调整。
        "Authorization": f"QiniuAKSK {QINIU_AK}:{QINIU_SK}",
    }
    if include_content_type:
        headers["Content-Type"] = "application/json"
    return headers


def _has_qiniu_text_credential() -> bool:
    return bool(QINIU_TEXT_API_KEY or (QINIU_AK and QINIU_SK))


def _has_qiniu_video_credential() -> bool:
    return bool(QINIU_VIDEO_API_KEY or QINIU_TEXT_API_KEY or (QINIU_AK and QINIU_SK))


def _build_qiniu_headers(scope: str, *, include_content_type: bool = True) -> Dict[str, str]:
    headers: Dict[str, str] = {}
    if include_content_type:
        headers["Content-Type"] = "application/json"

    # 优先走 sk-... API Key，和你当前网关报错信息保持一致。
    if scope == "text" and QINIU_TEXT_API_KEY:
        headers["Authorization"] = f"Bearer {QINIU_TEXT_API_KEY}"
        return headers
    if scope == "video" and QINIU_VIDEO_API_KEY:
        headers["Authorization"] = f"Bearer {QINIU_VIDEO_API_KEY}"
        return headers
    if scope == "video" and QINIU_TEXT_API_KEY:
        headers["Authorization"] = f"Bearer {QINIU_TEXT_API_KEY}"
        return headers

    return {**headers, **_build_qiniu_aksk_headers(include_content_type=False)}


def _get_model_candidates(primary_model: str) -> List[str]:
    candidates: List[str] = []
    if primary_model:
        candidates.append(primary_model)
    for item in QINIU_LLM_FALLBACK_MODELS:
        if item not in candidates:
            candidates.append(item)
    return candidates


def _call_provider_json(provider: str, system_prompt: str, user_prompt: str) -> Dict[str, Any]:
    provider_config = MODEL_PROVIDERS.get(provider)
    if not provider_config:
        raise ValueError(f"Unknown provider: {provider}")

    if provider_config.get("requires_qiniu_credential") and (not _has_qiniu_text_credential()):
        raise RuntimeError("QINIU_TEXT_API_KEY or QINIU_AK/QINIU_SK is missing in .env")
    
    return _call_openai_compatible_json(provider_config, system_prompt, user_prompt)


def _call_provider_text(provider: str, system_prompt: str, user_prompt: str) -> str:
    provider_config = MODEL_PROVIDERS.get(provider)
    if not provider_config:
        raise ValueError(f"Unknown provider: {provider}")

    if provider_config.get("requires_qiniu_credential") and (not _has_qiniu_text_credential()):
        raise RuntimeError("QINIU_TEXT_API_KEY or QINIU_AK/QINIU_SK is missing in .env")
    
    return _call_openai_compatible_text(provider_config, system_prompt, user_prompt)


def _call_openai_compatible_json(provider_config: Dict[str, Any], system_prompt: str, user_prompt: str) -> Dict[str, Any]:
    url = _resolve_url(provider_config.get("base_url", ""), "/chat/completions")
    headers: Dict[str, str] = {"Content-Type": "application/json"}
    if provider_config.get("type") == "qiniu_openai_compatible":
        headers.update(_build_qiniu_headers("text", include_content_type=False))
    else:
        api_key = str(provider_config.get("api_key", "")).strip()
        if not api_key:
            raise RuntimeError(f"{provider_config.get('name', 'Provider')} API key is missing in .env")
        headers["Authorization"] = f"Bearer {api_key}"
    model_candidates = _get_model_candidates(str(provider_config.get("model", "")).strip())
    if not model_candidates:
        raise RuntimeError("No model configured for qiniu provider")

    last_error: Optional[Exception] = None
    for model_name in model_candidates:
        body = {
            "model": model_name,
            "temperature": 0.7,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        }

        try:
            response = requests.post(url, headers=headers, json=body, timeout=60, verify=False)
            response.raise_for_status()
            data = response.json()
            content = data["choices"][0]["message"]["content"]
            return _extract_json(content)
        except requests.HTTPError as e:
            last_error = e
            detail = (e.response.text if e.response is not None else "").lower()
            # 七牛网关常见报错：模型无可用通道，则自动尝试下一个候选模型。
            if "no available channels for model" in detail:
                continue
            raise

    if last_error:
        raise last_error
    raise RuntimeError("Model request failed")


def _call_openai_compatible_text(provider_config: Dict[str, Any], system_prompt: str, user_prompt: str) -> str:
    url = _resolve_url(provider_config.get("base_url", ""), "/chat/completions")
    headers: Dict[str, str] = {"Content-Type": "application/json"}
    if provider_config.get("type") == "qiniu_openai_compatible":
        headers.update(_build_qiniu_headers("text", include_content_type=False))
    else:
        api_key = str(provider_config.get("api_key", "")).strip()
        if not api_key:
            raise RuntimeError(f"{provider_config.get('name', 'Provider')} API key is missing in .env")
        headers["Authorization"] = f"Bearer {api_key}"
    model_candidates = _get_model_candidates(str(provider_config.get("model", "")).strip())
    if not model_candidates:
        raise RuntimeError("No model configured for qiniu provider")

    last_error: Optional[Exception] = None
    for model_name in model_candidates:
        body = {
            "model": model_name,
            "temperature": 0.8,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        }

        try:
            response = requests.post(url, headers=headers, json=body, timeout=60, verify=False)
            response.raise_for_status()
            data = response.json()
            return data["choices"][0]["message"]["content"].strip()
        except requests.HTTPError as e:
            last_error = e
            detail = (e.response.text if e.response is not None else "").lower()
            if "no available channels for model" in detail:
                continue
            raise

    if last_error:
        raise last_error
    raise RuntimeError("Model request failed")


def _call_qwen_json(system_prompt: str, user_prompt: str) -> Dict[str, Any]:
    return _call_provider_json(DEFAULT_PROVIDER, system_prompt, user_prompt)


def _call_qwen_text(system_prompt: str, user_prompt: str) -> str:
    return _call_provider_text(DEFAULT_PROVIDER, system_prompt, user_prompt)


def _video_script_prompt(payload: Dict[str, Any]) -> str:
    return f"""
你是短剧编导，请直接输出一段可拍可生视频的短剧脚本。

输入信息：
- 题材：{payload.get('genre', '')}
- 核心设定：{payload.get('idea', '')}
- 人物：{payload.get('roles', '')}
- 风格：{payload.get('style', '')}
- 时长：{payload.get('duration_sec', 10)} 秒

输出要求：
1) 先给出“标题”。
2) 再给“短剧脚本（分镜级）”，包含 4-6 个镜头，每个镜头写画面、动作、台词/音效。
3) 最后给“视频生成提示词（中文）”，用于文生视频，保证画面连贯。
4) 全文中文，简洁有戏剧冲突。
""".strip()


def _looks_like_vidu_queue_path(path_text: str) -> bool:
    return "queue/fal-ai/vidu" in str(path_text or "").lower()


def _pick_vidu_create_path(payload: Dict[str, Any], model: str, configured_path: str) -> str:
    mode = str(payload.get("video_mode", "")).strip().lower()
    if mode in {"text", "text-to-video", "t2v"}:
        return QINIU_VIDU_Q3_TEXT_TO_VIDEO_PATH
    if mode in {"image", "image-to-video", "i2v"}:
        return QINIU_VIDU_Q3_IMAGE_TO_VIDEO_PATH
    if mode in {"start_end", "start-end", "start-end-to-video", "s2v"}:
        return QINIU_VIDU_Q3_START_END_TO_VIDEO_PATH

    model_text = str(model or "").lower()
    if "viduq3" not in model_text:
        return configured_path

    start_image_url = str(payload.get("start_image_url", "")).strip()
    end_image_url = str(payload.get("end_image_url", "")).strip()
    image_url = str(payload.get("image_url", "")).strip()

    if start_image_url and end_image_url:
        return QINIU_VIDU_Q3_START_END_TO_VIDEO_PATH
    if image_url:
        return QINIU_VIDU_Q3_IMAGE_TO_VIDEO_PATH
    return QINIU_VIDU_Q3_TEXT_TO_VIDEO_PATH


def _size_to_vidu_resolution(size: str) -> str:
    text = str(size or "").strip().lower().replace("x", "*")
    mapping = {
        "1920*1080": "1080p",
        "1280*720": "720p",
        "960*540": "540p",
    }
    return mapping.get(text, "720p")


def _map_task_status(raw_status: str) -> str:
    status = str(raw_status or "").upper()
    if status in {"SUCCEEDED", "FAILED", "CANCELED", "IN_PROGRESS", "IN_QUEUE", "PENDING"}:
        return status
    if status in {"COMPLETED", "DONE", "SUCCESS"}:
        return "SUCCEEDED"
    if status in {"QUEUED"}:
        return "IN_QUEUE"
    if status in {"RUNNING", "PROCESSING"}:
        return "IN_PROGRESS"
    if status in {"CANCELLED"}:
        return "CANCELED"
    return status or "UNKNOWN"


def _normalize_create_video_response(data: Dict[str, Any]) -> Dict[str, Any]:
    if not isinstance(data, dict):
        return {"output": {"task_id": "", "task_status": "UNKNOWN"}, "raw": data}

    if isinstance(data.get("output"), dict):
        return data

    request_id = str(data.get("request_id", "")).strip()
    status = _map_task_status(str(data.get("status", "")).strip())
    if request_id:
        return {
            "output": {
                "task_id": request_id,
                "task_status": status or "IN_QUEUE",
                "status_url": data.get("status_url", ""),
                "response_url": data.get("response_url", ""),
                "cancel_url": data.get("cancel_url", ""),
            },
            "raw": data,
        }

    return {"output": {"task_id": "", "task_status": status or "UNKNOWN"}, "raw": data}


def _normalize_query_video_response(data: Dict[str, Any], task_id: str) -> Dict[str, Any]:
    if not isinstance(data, dict):
        return {
            "output": {
                "task_id": task_id,
                "task_status": "UNKNOWN",
                "video_url": "",
            },
            "raw": data,
        }

    if isinstance(data.get("output"), dict):
        output = data.get("output", {})
        output = dict(output)
        output["task_id"] = output.get("task_id") or task_id
        output["task_status"] = _map_task_status(str(output.get("task_status", "")))
        return {**data, "output": output}

    status = _map_task_status(str(data.get("status", "")).strip())
    result = data.get("result") if isinstance(data.get("result"), dict) else {}
    video = result.get("video") if isinstance(result.get("video"), dict) else {}
    video_url = str(video.get("url", "") or "").strip()
    request_id = str(data.get("request_id", "") or task_id)

    return {
        "output": {
            "task_id": request_id,
            "task_status": status or "UNKNOWN",
            "video_url": video_url,
            "url": video_url,
            "status_url": data.get("status_url", ""),
            "response_url": data.get("response_url", ""),
            "cancel_url": data.get("cancel_url", ""),
        },
        "raw": data,
    }


def _create_video_task(payload: Dict[str, Any]) -> Dict[str, Any]:
    if not _has_qiniu_video_credential():
        raise RuntimeError("QINIU_VIDEO_API_KEY or QINIU_AK/QINIU_SK is missing in .env")

    prompt = str(payload.get("prompt", "")).strip()
    if not prompt:
        raise ValueError("Video prompt is required.")

    model = payload.get("model", QINIU_VIDEO_MODEL)
    size = payload.get("size", "1280*720")
    duration = int(payload.get("duration", 10))
    prompt_extend = bool(payload.get("prompt_extend", True))

    create_path = _pick_vidu_create_path(payload, str(model), QINIU_VIDEO_CREATE_PATH)
    is_vidu_queue = _looks_like_vidu_queue_path(create_path)

    headers = _build_qiniu_headers("video", include_content_type=True)
    if QINIU_ENABLE_ASYNC_HEADER:
        headers["X-DashScope-Async"] = "enable"

    if is_vidu_queue:
        body: Dict[str, Any] = {
            "prompt": prompt,
            "duration": duration,
            "resolution": str(payload.get("resolution", "")).strip() or _size_to_vidu_resolution(str(size)),
            "movement_amplitude": str(payload.get("movement_amplitude", "auto")).strip() or "auto",
            "seed": int(payload.get("seed", 0) or 0),
            "is_rec": bool(payload.get("is_rec", False)),
            "audio": bool(payload.get("audio", False)),
        }

        if str(payload.get("voice_id", "")).strip():
            body["voice_id"] = str(payload.get("voice_id", "")).strip()

        if "start-end-to-video" in str(create_path):
            start_image_url = str(payload.get("start_image_url", "")).strip()
            end_image_url = str(payload.get("end_image_url", "")).strip()
            if not start_image_url or not end_image_url:
                raise ValueError("start_image_url and end_image_url are required for start-end-to-video endpoint.")
            body["start_image_url"] = start_image_url
            body["end_image_url"] = end_image_url
        elif "text-to-video" in str(create_path):
            body["style"] = str(payload.get("style", "general")).strip() or "general"
            body["aspect_ratio"] = str(payload.get("aspect_ratio", "16:9")).strip() or "16:9"
            body["bgm"] = bool(payload.get("bgm", False))
        else:
            image_url = str(payload.get("image_url", "")).strip()
            if not image_url:
                raise ValueError("image_url is required for image-to-video endpoint.")
            body["image_url"] = image_url
    else:
        body = {
            "model": model,
            "input": {"prompt": prompt},
            "parameters": {
                "size": size,
                "duration": duration,
                "duration_seconds": duration,
                "durationSeconds": duration,
                "prompt_extend": prompt_extend,
                "watermark": False,
            },
        }

    last_error: Optional[requests.HTTPError] = None
    url_candidates = _build_url_candidates(QINIU_VIDEO_BASE_URL, create_path)

    for idx, url in enumerate(url_candidates):
        try:
            response = requests.post(url, headers=headers, json=body, timeout=60, verify=False)
            response.raise_for_status()
            return _normalize_create_video_response(response.json())
        except requests.HTTPError as err:
            last_error = err
            can_retry = idx < len(url_candidates) - 1 and _is_not_found_or_method_not_allowed(err)
            if can_retry:
                continue
            raise

    if last_error:
        raise last_error
    raise RuntimeError("Video task creation failed")


def _query_video_task(task_id: str) -> Dict[str, Any]:
    if not _has_qiniu_video_credential():
        raise RuntimeError("QINIU_VIDEO_API_KEY or QINIU_AK/QINIU_SK is missing in .env")

    task_path = QINIU_VIDEO_TASK_PATH_TEMPLATE.format(task_id=task_id)
    headers = _build_qiniu_headers("video", include_content_type=False)

    last_error: Optional[requests.HTTPError] = None
    url_candidates = _build_url_candidates(QINIU_VIDEO_BASE_URL, task_path)

    for idx, url in enumerate(url_candidates):
        try:
            response = requests.get(url, headers=headers, timeout=60, verify=False)
            response.raise_for_status()
            return _normalize_query_video_response(response.json(), task_id)
        except requests.HTTPError as err:
            last_error = err
            can_retry = idx < len(url_candidates) - 1 and _is_not_found_or_method_not_allowed(err)
            if can_retry:
                continue
            raise

    if last_error:
        raise last_error
    raise RuntimeError("Video task query failed")


def _qiniu_web_search(query: str, max_results: int = 10, search_type: str = "web") -> Dict[str, Any]:
    if not _has_qiniu_text_credential():
        raise RuntimeError("QINIU_TEXT_API_KEY (or OPENAI_API_KEY) is missing in .env")

    payload = {
        "query": query,
        "max_results": int(max_results),
        "search_type": search_type,
    }
    url = _resolve_url(QINIU_LLM_BASE_URL, QINIU_WEB_SEARCH_PATH)
    headers = _build_qiniu_headers("text", include_content_type=True)
    response = requests.post(url, headers=headers, json=payload, timeout=60, verify=False)
    response.raise_for_status()
    return response.json()


def _extend_video_prompts(
    total_duration: int,
    segment_duration: int,
    base_prompt: str,
    provider: str,
) -> List[Dict[str, Any]]:
    """Split long video into segments and let LLM draft each segment prompt.

    This works purely在文本层面：模型根据整体提示和前面各段的提示词，续写后续段落的描述，
    以便逐段调用万相文生视频接口。这里不会真正解析视频画面。
    """

    if total_duration <= 0:
        raise ValueError("total_duration must be positive")

    segment_duration = max(1, min(segment_duration, 10))
    segments: List[Dict[str, Any]] = []

    remaining = total_duration
    index = 1
    history: List[Dict[str, Any]] = []

    while remaining > 0:
        current_duration = min(segment_duration, remaining)

        if index == 1:
            segment_prompt = (
                f"【长视频第1段（0-{current_duration}秒）】在整体设定基础上，"
                f"生成这一段的详细画面描述，用于文生视频提示词。整体提示如下：\n{base_prompt}"
            )
        else:
            system_prompt = (
                "你是视频提示词续写助手。现在有一个长视频被拆成多段，每段约10秒。"\
                "你需要根据前面已经生成的各段提示词，总结已经发生的剧情，并继续写出下一段的提示词，"\
                "保证人物、场景、风格和镜头运动尽量连贯。输出一段可直接用于文生视频的中文提示词，不要解释。"
            )

            user_prompt = (
                "整体设定（用户原始提示）如下：\n"\
                f"{base_prompt}\n\n"\
                "前面已经规划好的各段提示词与时长：\n"\
                f"{json.dumps(history, ensure_ascii=False, indent=2)}\n\n"\
                f"现在请继续编写第{index}段（预计时长 {current_duration} 秒）的提示词，只输出这一段的提示词。"
            )

            segment_prompt = _call_provider_text(provider, system_prompt, user_prompt)

        segment = {
            "index": index,
            "duration": current_duration,
            "prompt": segment_prompt.strip(),
        }
        segments.append(segment)
        history.append({
            "index": index,
            "duration": current_duration,
            "prompt": segment_prompt.strip(),
        })

        remaining -= current_duration
        index += 1

    return segments


def _story_engine_prompt(payload: Dict[str, Any]) -> str:
    return f"""
你是短剧创作智能体的第一层：故事引擎。
请基于用户输入，输出严格 JSON（不要解释文字）。

用户输入：
- 创意: {payload.get('idea', '')}
- 主题偏好: {payload.get('theme', '')}
- 情绪基调: {payload.get('tone', '')}
- 结构模板偏好: {payload.get('structure', '')}

JSON schema:
{{
  "story_card": {{
    "logline": "一句话故事梗概",
    "theme": "核心主题",
    "tone": "情绪基调",
    "structure_template": "所用结构模板",
    "core_conflict": "核心冲突",
    "anchor_points": ["开端锚点", "转折锚点", "高潮锚点", "结局锚点"],
    "hook": "前三秒抓人钩子",
    "ending_type": "开放式/反转式/治愈式等"
  }},
  "next_questions": ["建议用户补充的问题1", "建议用户补充的问题2"]
}}
""".strip()


def _workshop_prompt(payload: Dict[str, Any]) -> str:
    story_card = payload.get("story_card", {})
    role_requirements = payload.get("role_requirements", "")
    plot_requirements = payload.get("plot_requirements", "")

    return f"""
You are the workshop layer of a short-drama writing assistant.
Return strict JSON only. No markdown, no explanation.
Write all content values in Chinese.

Story card:
{json.dumps(story_card, ensure_ascii=False, indent=2)}

Role requirements:
{role_requirements}

Plot requirements:
{plot_requirements}

Required JSON schema:
{{
  "characters": [
    {{
      "name": "角色名",
      "tags": ["职业", "性格", "目标", "缺陷"],
      "motivation": "核心动机",
      "arc": "角色弧光"
    }}
  ],
  "relationships": [
    {{
      "from": "角色A",
      "to": "角色B",
      "type": "关系类型",
      "tension": "冲突点"
    }}
  ],
  "plot_nodes": [
    {{
      "id": "N1",
      "template_stage": "激励事件/第一次转折/高潮等",
      "summary": "这一节点发生了什么",
      "location": "地点",
      "action_draft": "动作与场面调度",
      "dialogue_draft": ["角色: 台词"],
      "emotion_shift": "这一节点前后情绪变化",
      "consistency_check": "潜在逻辑问题，没有则写无"
    }}
  ],
  "timeline_view": ["N1", "N2"],
  "card_wall_groups": [
    {{
      "group": "铺垫/冲突/反转/收束",
      "node_ids": ["N1", "N2"]
    }}
  ]
}}
""".strip()


def _storyboard_prompt(payload: Dict[str, Any]) -> str:
    workshop = payload.get("workshop", {})
    style = payload.get("visual_style", "")

    return f"""
You are the storyboard layer of a short-drama writing assistant.
Return strict JSON only. No markdown, no explanation.
Write all content values in Chinese.

Workshop result:
{json.dumps(workshop, ensure_ascii=False, indent=2)}

Visual style requirement:
{style}

Required JSON schema:
{{
  "storyboards": [
    {{
      "shot_id": "S1",
      "related_node_id": "N1",
      "shot_type": "特写/中景/全景",
      "camera_movement": "固定/推/拉/摇/移/跟拍",
      "visual_description": "画面内容",
      "dialogue_or_sfx": "对白或音效",
      "duration_sec": 4,
      "shooting_note": "拍摄备注",
      "prompt_draft": "可直接用于视频生成的提示词"
    }}
  ],
  "estimated_total_duration_sec": 60,
  "export_ready_checklist": ["服化道", "场景", "收音", "灯光"]
}}
""".strip()


def _command_prompt(payload: Dict[str, Any]) -> str:
    command = payload.get("command", "")
    project_state = payload.get("project_state", {})

    return f"""
你是短剧创作助手的全局指令执行器。
请读取当前状态并执行用户自然语言命令，输出严格 JSON。

用户命令：{command}
当前项目状态：
{json.dumps(project_state, ensure_ascii=False, indent=2)}

JSON schema:
{{
  "command_understanding": "你对命令的理解",
  "updated_state": {{
    "story_card": {{}} ,
    "workshop": {{}} ,
    "storyboard": {{}}
  }},
  "consistency_report": ["一致性检查结果"],
  "suggestions": ["下一步建议"]
}}
""".strip()


def _normalize_export_payload(payload: Dict[str, Any]) -> Dict[str, Any]:
    project = payload.get("project", {}) if isinstance(payload, dict) else {}
    if not isinstance(project, dict):
        project = {}

    return {
        "project": {
            "id": _as_text(project.get("id")),
            "name": _as_text(_first_present(project, "name", "project_name")) or "未命名项目",
            "creator": _as_text(project.get("creator")),
            "description": _as_text(project.get("description")),
            "created_at": _as_text(project.get("created_at")),
            "updated_at": _as_text(project.get("updated_at")),
        },
        "current_provider": _as_text(
            payload.get("current_provider") if isinstance(payload, dict) else ""
        )
        or _as_text(payload.get("provider") if isinstance(payload, dict) else "")
        or _as_text(project.get("last_provider")),
        "exported_at": _as_text(payload.get("exported_at") if isinstance(payload, dict) else "") or _utc_now_iso(),
        "story_card": _normalize_story_card(payload.get("story_card") if isinstance(payload, dict) else None),
        "workshop": _normalize_workshop_result(payload.get("workshop") if isinstance(payload, dict) else None),
        "storyboard": _normalize_storyboard_result(payload.get("storyboard") if isinstance(payload, dict) else None),
        "video_lab": _normalize_video_lab_state(payload.get("video_lab") if isinstance(payload, dict) else None),
    }


def _export_markdown(payload: Dict[str, Any]) -> Dict[str, Any]:
    data = _normalize_export_payload(payload)
    project = data["project"]
    story_card = data["story_card"] or {}
    workshop = data["workshop"] or {}
    storyboard = data["storyboard"] or {}
    video_lab = data["video_lab"] or _default_project_state()["video_lab"]

    lines: List[str] = []
    lines.append("# AI短剧项目导出")
    lines.append("")
    lines.append("## 0. 项目信息")
    lines.append(f"- 项目名称: {project.get('name', '未命名项目')}")
    lines.append(f"- 创建人: {project.get('creator', '') or '-'}")
    lines.append(f"- 项目描述: {project.get('description', '') or '-'}")
    lines.append(f"- 创建时间: {project.get('created_at', '') or '-'}")
    lines.append(f"- 更新时间: {project.get('updated_at', '') or '-'}")
    lines.append(f"- 导出时间: {data.get('exported_at', '') or '-'}")
    lines.append(f"- 当前模型: {data.get('current_provider', '') or '-'}")
    lines.append("")

    lines.append("## 1. 故事卡")
    lines.append(f"- Logline: {story_card.get('logline', '') or '-'}")
    lines.append(f"- 主题: {story_card.get('theme', '') or '-'}")
    lines.append(f"- 基调: {story_card.get('tone', '') or '-'}")
    lines.append(f"- 结构模板: {story_card.get('structure_template', '') or '-'}")
    lines.append(f"- 核心冲突: {story_card.get('core_conflict', '') or '-'}")
    lines.append(f"- 钩子: {story_card.get('hook', '') or '-'}")
    lines.append(f"- 结局类型: {story_card.get('ending_type', '') or '-'}")
    lines.append("- 结构锚点:")
    if story_card.get("anchor_points"):
        for index, point in enumerate(story_card.get("anchor_points", []), start=1):
            lines.append(f"  {index}. {point}")
    else:
        lines.append("  - 无")
    lines.append("")

    lines.append("## 2. 角色设定")
    if workshop.get("characters"):
        for character in workshop.get("characters", []):
            lines.append(
                f"- {character.get('name', '未命名角色')} | 标签: {', '.join(character.get('tags', [])) or '-'} | 动机: {character.get('motivation', '') or '-'} | 弧光: {character.get('arc', '') or '-'}"
            )
    else:
        lines.append("- 无")
    lines.append("")

    lines.append("## 3. 角色关系")
    if workshop.get("relationships"):
        for relationship in workshop.get("relationships", []):
            lines.append(
                f"- {relationship.get('from', '-')} -> {relationship.get('to', '-')} | 类型: {relationship.get('type', '') or '-'} | 冲突: {relationship.get('tension', '') or '-'}"
            )
    else:
        lines.append("- 无")
    lines.append("")

    lines.append("## 4. 剧情节点")
    if workshop.get("plot_nodes"):
        for node in workshop.get("plot_nodes", []):
            lines.append(f"- {node.get('id', '') or '-'} [{node.get('template_stage', '') or '-'}]")
            lines.append(f"  摘要: {node.get('summary', '') or '-'}")
            lines.append(f"  地点: {node.get('location', '') or '-'}")
            lines.append(f"  动作: {node.get('action_draft', '') or '-'}")
            lines.append(f"  对白: {_dialogue_text(node.get('dialogue_draft'))}")
            lines.append(f"  情绪变化: {node.get('emotion_shift', '') or '-'}")
            lines.append(f"  一致性检查: {node.get('consistency_check', '') or '-'}")
    else:
        lines.append("- 无")
    lines.append("")

    lines.append("## 5. 时间线顺序")
    if workshop.get("timeline_view"):
        for index, node_id in enumerate(workshop.get("timeline_view", []), start=1):
            lines.append(f"{index}. {node_id}")
    else:
        lines.append("- 无")
    lines.append("")

    lines.append("## 6. 分镜表")
    lines.append("| 镜头ID | 关联节点 | 景别 | 运镜 | 时长(秒) | 画面 | 对白/音效 | 提示词 | 备注 |")
    lines.append("|---|---|---|---|---|---|---|---|---|")
    if storyboard.get("storyboards"):
        for shot in storyboard.get("storyboards", []):
            lines.append(
                "| {shot_id} | {node} | {shot_type} | {move} | {duration} | {visual} | {sound} | {prompt} | {note} |".format(
                    shot_id=shot.get("shot_id", ""),
                    node=shot.get("related_node_id", ""),
                    shot_type=shot.get("shot_type", ""),
                    move=shot.get("camera_movement", ""),
                    duration=shot.get("duration_sec", ""),
                    visual=str(shot.get("visual_description", "")).replace("|", "\\|"),
                    sound=str(shot.get("dialogue_or_sfx", "")).replace("|", "\\|"),
                    prompt=str(shot.get("prompt_draft", "")).replace("|", "\\|"),
                    note=str(shot.get("shooting_note", "")).replace("|", "\\|"),
                )
            )
    lines.append("")
    lines.append(f"- 预估总时长: {storyboard.get('estimated_total_duration_sec', 0)} 秒")
    lines.append("- 导出前检查清单:")
    if storyboard.get("export_ready_checklist"):
        for item in storyboard.get("export_ready_checklist", []):
            lines.append(f"  - {item}")
    else:
        lines.append("  - 无")
    lines.append("")

    lines.append("## 7. 视频任务摘要")
    lines.append(f"- 当前任务ID: {video_lab.get('task_id', '') or '-'}")
    lines.append(f"- 当前任务状态: {video_lab.get('task_status', '') or '-'}")
    lines.append(f"- 当前视频链接: {video_lab.get('video_url', '') or '-'}")
    lines.append(f"- 长视频总时长: {video_lab.get('total_duration', 0)} 秒")
    lines.append(f"- 文件名前缀: {video_lab.get('filename_prefix', '') or '-'}")
    if video_lab.get("long_segments"):
        lines.append("- 拆段任务:")
        for segment in video_lab.get("long_segments", []):
            lines.append(
                f"  - 第{segment.get('index', '-')}段 | 时长: {segment.get('duration', 0)} 秒 | Task ID: {segment.get('task_id', '') or '-'} | 状态: {segment.get('task_status', '') or '-'}"
            )
    else:
        lines.append("- 拆段任务: 无")

    return {"markdown": "\n".join(lines)}


def _build_docx(payload: Dict[str, Any]) -> BytesIO:
    data = _normalize_export_payload(payload)
    project = data["project"]
    story_card = data["story_card"] or {}
    workshop = data["workshop"] or {}
    storyboard = data["storyboard"] or {}
    video_lab = data["video_lab"] or _default_project_state()["video_lab"]

    doc = Document()
    doc.add_heading("AI短剧项目导出", level=1)

    doc.add_heading("0. 项目信息", level=2)
    for line in [
        f"项目名称: {project.get('name', '未命名项目')}",
        f"创建人: {project.get('creator', '') or '-'}",
        f"项目描述: {project.get('description', '') or '-'}",
        f"创建时间: {project.get('created_at', '') or '-'}",
        f"更新时间: {project.get('updated_at', '') or '-'}",
        f"导出时间: {data.get('exported_at', '') or '-'}",
        f"当前模型: {data.get('current_provider', '') or '-'}",
    ]:
        doc.add_paragraph(line)

    doc.add_heading("1. 故事卡", level=2)
    for line in [
        f"Logline: {story_card.get('logline', '') or '-'}",
        f"主题: {story_card.get('theme', '') or '-'}",
        f"基调: {story_card.get('tone', '') or '-'}",
        f"结构模板: {story_card.get('structure_template', '') or '-'}",
        f"核心冲突: {story_card.get('core_conflict', '') or '-'}",
        f"钩子: {story_card.get('hook', '') or '-'}",
        f"结局类型: {story_card.get('ending_type', '') or '-'}",
    ]:
        doc.add_paragraph(line)
    doc.add_paragraph("结构锚点:")
    if story_card.get("anchor_points"):
        for point in story_card.get("anchor_points", []):
            doc.add_paragraph(point, style="List Bullet")
    else:
        doc.add_paragraph("无", style="List Bullet")

    doc.add_heading("2. 角色设定", level=2)
    if workshop.get("characters"):
        for character in workshop.get("characters", []):
            doc.add_paragraph(
                f"{character.get('name', '未命名角色')} | 标签: {', '.join(character.get('tags', [])) or '-'} | 动机: {character.get('motivation', '') or '-'} | 弧光: {character.get('arc', '') or '-'}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("无", style="List Bullet")

    doc.add_heading("3. 角色关系", level=2)
    if workshop.get("relationships"):
        for relationship in workshop.get("relationships", []):
            doc.add_paragraph(
                f"{relationship.get('from', '-')} -> {relationship.get('to', '-')} | 类型: {relationship.get('type', '') or '-'} | 冲突: {relationship.get('tension', '') or '-'}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("无", style="List Bullet")

    doc.add_heading("4. 剧情节点", level=2)
    if workshop.get("plot_nodes"):
        for node in workshop.get("plot_nodes", []):
            doc.add_paragraph(f"{node.get('id', '')} [{node.get('template_stage', '')}]", style="List Bullet")
            for detail in [
                f"摘要: {node.get('summary', '') or '-'}",
                f"地点: {node.get('location', '') or '-'}",
                f"动作: {node.get('action_draft', '') or '-'}",
                f"对白: {_dialogue_text(node.get('dialogue_draft'))}",
                f"情绪变化: {node.get('emotion_shift', '') or '-'}",
                f"一致性检查: {node.get('consistency_check', '') or '-'}",
            ]:
                doc.add_paragraph(detail)
    else:
        doc.add_paragraph("无", style="List Bullet")

    doc.add_heading("5. 时间线顺序", level=2)
    if workshop.get("timeline_view"):
        for node_id in workshop.get("timeline_view", []):
            doc.add_paragraph(node_id, style="List Number")
    else:
        doc.add_paragraph("无", style="List Bullet")

    doc.add_heading("6. 分镜表", level=2)
    table = doc.add_table(rows=1, cols=9)
    headers = ["镜头ID", "关联节点", "景别", "运镜", "时长", "画面", "对白/音效", "提示词", "备注"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for shot in storyboard.get("storyboards", []):
        row = table.add_row().cells
        row[0].text = str(shot.get("shot_id", ""))
        row[1].text = str(shot.get("related_node_id", ""))
        row[2].text = str(shot.get("shot_type", ""))
        row[3].text = str(shot.get("camera_movement", ""))
        row[4].text = str(shot.get("duration_sec", ""))
        row[5].text = str(shot.get("visual_description", ""))
        row[6].text = str(shot.get("dialogue_or_sfx", ""))
        row[7].text = str(shot.get("prompt_draft", ""))
        row[8].text = str(shot.get("shooting_note", ""))
    doc.add_paragraph(f"预估总时长: {storyboard.get('estimated_total_duration_sec', 0)} 秒")
    doc.add_paragraph(
        "导出前检查清单: "
        + (", ".join(storyboard.get("export_ready_checklist", [])) if storyboard.get("export_ready_checklist") else "无")
    )

    doc.add_heading("7. 视频任务摘要", level=2)
    for line in [
        f"当前任务ID: {video_lab.get('task_id', '') or '-'}",
        f"当前任务状态: {video_lab.get('task_status', '') or '-'}",
        f"当前视频链接: {video_lab.get('video_url', '') or '-'}",
        f"长视频总时长: {video_lab.get('total_duration', 0)} 秒",
        f"文件名前缀: {video_lab.get('filename_prefix', '') or '-'}",
    ]:
        doc.add_paragraph(line)
    if video_lab.get("long_segments"):
        for segment in video_lab.get("long_segments", []):
            doc.add_paragraph(
                f"第{segment.get('index', '-')}段 | 时长: {segment.get('duration', 0)} 秒 | Task ID: {segment.get('task_id', '') or '-'} | 状态: {segment.get('task_status', '') or '-'}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("拆段任务: 无", style="List Bullet")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _build_pdf(payload: Dict[str, Any]) -> BytesIO:
    data = _normalize_export_payload(payload)
    project = data["project"]
    story_card = data["story_card"] or {}
    workshop = data["workshop"] or {}
    storyboard = data["storyboard"] or {}
    video_lab = data["video_lab"] or _default_project_state()["video_lab"]

    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    font_name = _register_pdf_font()

    def ensure_space(y_pos: float, need: float = 40.0) -> float:
        if y_pos < need:
            c.showPage()
            c.setFont(font_name, 11)
            return height - 50
        return y_pos

    def draw_section(title: str, lines: List[str], y_pos: float) -> float:
        y_pos = ensure_space(y_pos, 60)
        c.setFont(font_name, 13)
        c.drawString(40, y_pos, title)
        y_pos -= 20
        c.setFont(font_name, 11)
        for line in lines:
            y_pos = ensure_space(y_pos)
            y_pos = _draw_wrapped(c, line, font_name, 11, 50, y_pos, width - 90)
        return y_pos - 8

    y = height - 50
    c.setFont(font_name, 16)
    c.drawString(40, y, "AI短剧项目导出")
    y -= 30

    y = draw_section(
        "0. 项目信息",
        [
            f"项目名称: {project.get('name', '未命名项目')}",
            f"创建人: {project.get('creator', '') or '-'}",
            f"项目描述: {project.get('description', '') or '-'}",
            f"创建时间: {project.get('created_at', '') or '-'}",
            f"更新时间: {project.get('updated_at', '') or '-'}",
            f"导出时间: {data.get('exported_at', '') or '-'}",
            f"当前模型: {data.get('current_provider', '') or '-'}",
        ],
        y,
    )

    y = draw_section(
        "1. 故事卡",
        [
            f"Logline: {story_card.get('logline', '') or '-'}",
            f"主题: {story_card.get('theme', '') or '-'}",
            f"基调: {story_card.get('tone', '') or '-'}",
            f"结构模板: {story_card.get('structure_template', '') or '-'}",
            f"核心冲突: {story_card.get('core_conflict', '') or '-'}",
            f"钩子: {story_card.get('hook', '') or '-'}",
            f"结局类型: {story_card.get('ending_type', '') or '-'}",
            "结构锚点: " + (", ".join(story_card.get("anchor_points", [])) if story_card.get("anchor_points") else "无"),
        ],
        y,
    )

    role_lines = [
        f"- {character.get('name', '未命名角色')} | 标签: {', '.join(character.get('tags', [])) or '-'} | 动机: {character.get('motivation', '') or '-'} | 弧光: {character.get('arc', '') or '-'}"
        for character in workshop.get("characters", [])
    ] or ["- 无"]
    y = draw_section("2. 角色设定", role_lines, y)

    relation_lines = [
        f"- {relationship.get('from', '-')} -> {relationship.get('to', '-')} | 类型: {relationship.get('type', '') or '-'} | 冲突: {relationship.get('tension', '') or '-'}"
        for relationship in workshop.get("relationships", [])
    ] or ["- 无"]
    y = draw_section("3. 角色关系", relation_lines, y)

    plot_lines = []
    for node in workshop.get("plot_nodes", []):
        plot_lines.extend(
            [
                f"- {node.get('id', '')} [{node.get('template_stage', '')}]",
                f"  摘要: {node.get('summary', '') or '-'}",
                f"  地点: {node.get('location', '') or '-'}",
                f"  动作: {node.get('action_draft', '') or '-'}",
                f"  对白: {_dialogue_text(node.get('dialogue_draft'))}",
                f"  情绪变化: {node.get('emotion_shift', '') or '-'}",
                f"  一致性检查: {node.get('consistency_check', '') or '-'}",
            ]
        )
    if not plot_lines:
        plot_lines = ["- 无"]
    y = draw_section("4. 剧情节点", plot_lines, y)

    timeline_lines = [f"{index}. {node_id}" for index, node_id in enumerate(workshop.get("timeline_view", []), start=1)] or ["- 无"]
    y = draw_section("5. 时间线顺序", timeline_lines, y)

    storyboard_lines = []
    for shot in storyboard.get("storyboards", []):
        storyboard_lines.extend(
            [
                f"{shot.get('shot_id', '')} | 节点: {shot.get('related_node_id', '')} | 景别: {shot.get('shot_type', '')} | 运镜: {shot.get('camera_movement', '')} | 时长: {shot.get('duration_sec', 0)}秒",
                f"画面: {shot.get('visual_description', '') or '-'}",
                f"对白/音效: {shot.get('dialogue_or_sfx', '') or '-'}",
                f"提示词: {shot.get('prompt_draft', '') or '-'}",
                f"备注: {shot.get('shooting_note', '') or '-'}",
            ]
        )
    if not storyboard_lines:
        storyboard_lines = ["- 无"]
    storyboard_lines.append(f"预估总时长: {storyboard.get('estimated_total_duration_sec', 0)} 秒")
    storyboard_lines.append(
        "检查清单: "
        + (", ".join(storyboard.get("export_ready_checklist", [])) if storyboard.get("export_ready_checklist") else "无")
    )
    y = draw_section("6. 分镜表", storyboard_lines, y)

    video_lines = [
        f"当前任务ID: {video_lab.get('task_id', '') or '-'}",
        f"当前任务状态: {video_lab.get('task_status', '') or '-'}",
        f"当前视频链接: {video_lab.get('video_url', '') or '-'}",
        f"长视频总时长: {video_lab.get('total_duration', 0)} 秒",
        f"文件名前缀: {video_lab.get('filename_prefix', '') or '-'}",
    ]
    if video_lab.get("long_segments"):
        for segment in video_lab.get("long_segments", []):
            video_lines.append(
                f"第{segment.get('index', '-')}段 | 时长: {segment.get('duration', 0)} 秒 | Task ID: {segment.get('task_id', '') or '-'} | 状态: {segment.get('task_status', '') or '-'}"
            )
    else:
        video_lines.append("拆段任务: 无")
    y = draw_section("7. 视频任务摘要", video_lines, y)

    c.save()
    buffer.seek(0)
    return buffer


def _dialogue_text(dialogue: Any) -> str:
    values = _string_list(dialogue)
    return " / ".join(values) if values else "-"


def _normalize_export_payload(payload: Dict[str, Any]) -> Dict[str, Any]:
    project = payload.get("project", {}) if isinstance(payload, dict) else {}
    if not isinstance(project, dict):
        project = {}

    return {
        "project": {
            "id": _as_text(project.get("id")),
            "name": _as_text(_first_present(project, "name", "project_name")) or "\u672a\u547d\u540d\u9879\u76ee",
            "creator": _as_text(project.get("creator")),
            "description": _as_text(project.get("description")),
            "created_at": _as_text(project.get("created_at")),
            "updated_at": _as_text(project.get("updated_at")),
        },
        "current_provider": _as_text(payload.get("current_provider") if isinstance(payload, dict) else "")
        or _as_text(payload.get("provider") if isinstance(payload, dict) else "")
        or _as_text(project.get("last_provider")),
        "exported_at": _as_text(payload.get("exported_at") if isinstance(payload, dict) else "") or _utc_now_iso(),
        "story_card": _normalize_story_card(payload.get("story_card") if isinstance(payload, dict) else None),
        "workshop": _normalize_workshop_result(payload.get("workshop") if isinstance(payload, dict) else None),
        "storyboard": _normalize_storyboard_result(payload.get("storyboard") if isinstance(payload, dict) else None),
        "video_lab": _normalize_video_lab_state(payload.get("video_lab") if isinstance(payload, dict) else None),
    }


def _export_markdown(payload: Dict[str, Any]) -> Dict[str, Any]:
    data = _normalize_export_payload(payload)
    project = data["project"]
    story_card = data["story_card"] or {}
    workshop = data["workshop"] or {}
    storyboard = data["storyboard"] or {}
    video_lab = data["video_lab"] or _default_project_state()["video_lab"]

    lines: List[str] = []
    lines.append("# AI\u77ed\u5267\u9879\u76ee\u5bfc\u51fa")
    lines.append("")
    lines.append("## 0. \u9879\u76ee\u4fe1\u606f")
    lines.append(f"- \u9879\u76ee\u540d\u79f0: {project.get('name') or '未命名项目'}")
    lines.append(f"- \u521b\u5efa\u4eba: {project.get('creator', '') or '-'}")
    lines.append(f"- \u9879\u76ee\u63cf\u8ff0: {project.get('description', '') or '-'}")
    lines.append(f"- \u521b\u5efa\u65f6\u95f4: {project.get('created_at', '') or '-'}")
    lines.append(f"- \u66f4\u65b0\u65f6\u95f4: {project.get('updated_at', '') or '-'}")
    lines.append(f"- \u5bfc\u51fa\u65f6\u95f4: {data.get('exported_at', '') or '-'}")
    lines.append(f"- \u5f53\u524d\u6a21\u578b: {data.get('current_provider', '') or '-'}")
    lines.append("")

    lines.append("## 1. \u6545\u4e8b\u5361")
    lines.append(f"- Logline: {story_card.get('logline', '') or '-'}")
    lines.append(f"- \u4e3b\u9898: {story_card.get('theme', '') or '-'}")
    lines.append(f"- \u57fa\u8c03: {story_card.get('tone', '') or '-'}")
    lines.append(f"- \u7ed3\u6784\u6a21\u677f: {story_card.get('structure_template', '') or '-'}")
    lines.append(f"- \u6838\u5fc3\u51b2\u7a81: {story_card.get('core_conflict', '') or '-'}")
    lines.append(f"- \u94a9\u5b50: {story_card.get('hook', '') or '-'}")
    lines.append(f"- \u7ed3\u5c40\u7c7b\u578b: {story_card.get('ending_type', '') or '-'}")
    lines.append("- \u7ed3\u6784\u951a\u70b9:")
    if story_card.get("anchor_points"):
        for index, point in enumerate(story_card.get("anchor_points", []), start=1):
            lines.append(f"  {index}. {point}")
    else:
        lines.append("  - \u65e0")
    lines.append("")

    lines.append("## 2. \u89d2\u8272\u8bbe\u5b9a")
    if workshop.get("characters"):
        for character in workshop.get("characters", []):
            lines.append(
                f"- {character.get('name') or '未命名角色'} | \u6807\u7b7e: {', '.join(character.get('tags', [])) or '-'} | \u52a8\u673a: {character.get('motivation', '') or '-'} | \u5f27\u5149: {character.get('arc', '') or '-'}"
            )
    else:
        lines.append("- \u65e0")
    lines.append("")

    lines.append("## 3. \u89d2\u8272\u5173\u7cfb")
    if workshop.get("relationships"):
        for relationship in workshop.get("relationships", []):
            lines.append(
                f"- {relationship.get('from', '-')} -> {relationship.get('to', '-')} | \u7c7b\u578b: {relationship.get('type', '') or '-'} | \u51b2\u7a81: {relationship.get('tension', '') or '-'}"
            )
    else:
        lines.append("- \u65e0")
    lines.append("")

    lines.append("## 4. \u5267\u60c5\u8282\u70b9")
    if workshop.get("plot_nodes"):
        for node in workshop.get("plot_nodes", []):
            lines.append(f"- {node.get('id', '') or '-'} [{node.get('template_stage', '') or '-'}]")
            lines.append(f"  \u6458\u8981: {node.get('summary', '') or '-'}")
            lines.append(f"  \u5730\u70b9: {node.get('location', '') or '-'}")
            lines.append(f"  \u52a8\u4f5c: {node.get('action_draft', '') or '-'}")
            lines.append(f"  \u5bf9\u767d: {_dialogue_text(node.get('dialogue_draft'))}")
            lines.append(f"  \u60c5\u611f\u53d8\u5316: {node.get('emotion_shift', '') or '-'}")
            lines.append(f"  \u4e00\u81f4\u6027\u68c0\u67e5: {node.get('consistency_check', '') or '-'}")
    else:
        lines.append("- \u65e0")
    lines.append("")

    lines.append("## 5. \u65f6\u95f4\u7ebf\u987a\u5e8f")
    if workshop.get("timeline_view"):
        for index, node_id in enumerate(workshop.get("timeline_view", []), start=1):
            lines.append(f"{index}. {node_id}")
    else:
        lines.append("- \u65e0")
    lines.append("")

    lines.append("## 6. \u5206\u955c\u8868")
    lines.append("| \u955c\u5934ID | \u5173\u8054\u8282\u70b9 | \u666f\u522b | \u8fd0\u955c | \u65f6\u957f(\u79d2) | \u753b\u9762 | \u5bf9\u767d/\u97f3\u6548 | \u63d0\u793a\u8bcd | \u5907\u6ce8 |")
    lines.append("|---|---|---|---|---|---|---|---|---|")
    if storyboard.get("storyboards"):
        for shot in storyboard.get("storyboards", []):
            lines.append(
                "| {shot_id} | {node} | {shot_type} | {move} | {duration} | {visual} | {sound} | {prompt} | {note} |".format(
                    shot_id=shot.get("shot_id", ""),
                    node=shot.get("related_node_id", ""),
                    shot_type=shot.get("shot_type", ""),
                    move=shot.get("camera_movement", ""),
                    duration=shot.get("duration_sec", ""),
                    visual=str(shot.get("visual_description", "")).replace("|", "\\|"),
                    sound=str(shot.get("dialogue_or_sfx", "")).replace("|", "\\|"),
                    prompt=str(shot.get("prompt_draft", "")).replace("|", "\\|"),
                    note=str(shot.get("shooting_note", "")).replace("|", "\\|"),
                )
            )
    lines.append("")
    lines.append(f"- \u9884\u4f30\u603b\u65f6\u957f: {storyboard.get('estimated_total_duration_sec', 0)} \u79d2")
    lines.append("- \u5bfc\u51fa\u524d\u68c0\u67e5\u6e05\u5355:")
    if storyboard.get("export_ready_checklist"):
        for item in storyboard.get("export_ready_checklist", []):
            lines.append(f"  - {item}")
    else:
        lines.append("  - \u65e0")
    lines.append("")

    lines.append("## 7. \u89c6\u9891\u4efb\u52a1\u6458\u8981")
    lines.append(f"- \u5f53\u524d\u4efb\u52a1ID: {video_lab.get('task_id', '') or '-'}")
    lines.append(f"- \u5f53\u524d\u4efb\u52a1\u72b6\u6001: {video_lab.get('task_status', '') or '-'}")
    lines.append(f"- \u5f53\u524d\u89c6\u9891\u94fe\u63a5: {video_lab.get('video_url', '') or '-'}")
    lines.append(f"- \u957f\u89c6\u9891\u603b\u65f6\u957f: {video_lab.get('total_duration', 0)} \u79d2")
    lines.append(f"- \u6587\u4ef6\u540d\u524d\u7f00: {video_lab.get('filename_prefix', '') or '-'}")
    if video_lab.get("long_segments"):
        lines.append("- \u62c6\u6bb5\u4efb\u52a1:")
        for segment in video_lab.get("long_segments", []):
            lines.append(
                f"  - \u7b2c{segment.get('index', '-')}段 | \u65f6\u957f: {segment.get('duration', 0)} \u79d2 | Task ID: {segment.get('task_id', '') or '-'} | \u72b6\u6001: {segment.get('task_status', '') or '-'}"
            )
    else:
        lines.append("- \u62c6\u6bb5\u4efb\u52a1: \u65e0")

    return {"markdown": "\n".join(lines)}


def _build_docx(payload: Dict[str, Any]) -> BytesIO:
    data = _normalize_export_payload(payload)
    project = data["project"]
    story_card = data["story_card"] or {}
    workshop = data["workshop"] or {}
    storyboard = data["storyboard"] or {}
    video_lab = data["video_lab"] or _default_project_state()["video_lab"]

    doc = Document()
    doc.add_heading("\u77ed\u5267\u9879\u76ee\u5bfc\u51fa", level=1)

    doc.add_heading("\u9879\u76ee\u4fe1\u606f", level=2)
    for line in [
        f"\u9879\u76ee\u540d\u79f0: {project.get('name') or '未命名项目'}",
        f"\u521b\u5efa\u4eba: {project.get('creator', '') or '-'}",
        f"\u9879\u76ee\u63cf\u8ff0: {project.get('description', '') or '-'}",
        f"\u521b\u5efa\u65f6\u95f4: {project.get('created_at', '') or '-'}",
        f"\u66f4\u65b0\u65f6\u95f4: {project.get('updated_at', '') or '-'}",
        f"\u5bfc\u51fa\u65f6\u95f4: {data.get('exported_at', '') or '-'}",
        f"\u5f53\u524d\u6a21\u578b: {data.get('current_provider', '') or '-'}",
    ]:
        doc.add_paragraph(line)

    doc.add_heading("\u6545\u4e8b\u5361", level=2)
    for line in [
        f"Logline: {story_card.get('logline', '') or '-'}",
        f"\u4e3b\u9898: {story_card.get('theme', '') or '-'}",
        f"\u57fa\u8c03: {story_card.get('tone', '') or '-'}",
        f"\u7ed3\u6784\u6a21\u677f: {story_card.get('structure_template', '') or '-'}",
        f"\u6838\u5fc3\u51b2\u7a81: {story_card.get('core_conflict', '') or '-'}",
        f"\u94a9\u5b50: {story_card.get('hook', '') or '-'}",
        f"\u7ed3\u5c40\u7c7b\u578b: {story_card.get('ending_type', '') or '-'}",
    ]:
        doc.add_paragraph(line)
    doc.add_paragraph("\u7ed3\u6784\u951a\u70b9:")
    if story_card.get("anchor_points"):
        for point in story_card.get("anchor_points", []):
            doc.add_paragraph(point, style="List Bullet")
    else:
        doc.add_paragraph("\u65e0", style="List Bullet")

    doc.add_heading("\u89d2\u8272\u8bbe\u5b9a", level=2)
    if workshop.get("characters"):
        for character in workshop.get("characters", []):
            doc.add_paragraph(
                f"{character.get('name') or '未命名角色'} | \u6807\u7b7e: {', '.join(character.get('tags', [])) or '-'} | \u52a8\u673a: {character.get('motivation', '') or '-'} | \u5f27\u5149: {character.get('arc', '') or '-'}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("\u65e0", style="List Bullet")

    doc.add_heading("\u89d2\u8272\u5173\u7cfb", level=2)
    if workshop.get("relationships"):
        for relationship in workshop.get("relationships", []):
            doc.add_paragraph(
                f"{relationship.get('from', '-')} -> {relationship.get('to', '-')} | \u7c7b\u578b: {relationship.get('type', '') or '-'} | \u51b2\u7a81: {relationship.get('tension', '') or '-'}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("\u65e0", style="List Bullet")

    doc.add_heading("\u5267\u60c5\u8282\u70b9", level=2)
    if workshop.get("plot_nodes"):
        for node in workshop.get("plot_nodes", []):
            doc.add_paragraph(f"{node.get('id', '')} [{node.get('template_stage', '')}]", style="List Bullet")
            for detail in [
                f"\u6458\u8981: {node.get('summary', '') or '-'}",
                f"\u5730\u70b9: {node.get('location', '') or '-'}",
                f"\u52a8\u4f5c: {node.get('action_draft', '') or '-'}",
                f"\u5bf9\u767d: {_dialogue_text(node.get('dialogue_draft'))}",
                f"\u60c5\u611f\u53d8\u5316: {node.get('emotion_shift', '') or '-'}",
                f"\u4e00\u81f4\u6027\u68c0\u67e5: {node.get('consistency_check', '') or '-'}",
            ]:
                doc.add_paragraph(detail)
    else:
        doc.add_paragraph("\u65e0", style="List Bullet")

    doc.add_heading("\u65f6\u95f4\u7ebf\u987a\u5e8f", level=2)
    if workshop.get("timeline_view"):
        for node_id in workshop.get("timeline_view", []):
            doc.add_paragraph(node_id, style="List Number")
    else:
        doc.add_paragraph("\u65e0", style="List Bullet")

    doc.add_heading("\u5206\u955c\u8868", level=2)
    table = doc.add_table(rows=1, cols=9)
    headers = ["\u955c\u5934ID", "\u5173\u8054\u8282\u70b9", "\u666f\u522b", "\u8fd0\u955c", "\u65f6\u957f", "\u753b\u9762", "\u5bf9\u767d/\u97f3\u6548", "\u63d0\u793a\u8bcd", "\u5907\u6ce8"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for shot in storyboard.get("storyboards", []):
        row = table.add_row().cells
        row[0].text = str(shot.get("shot_id", ""))
        row[1].text = str(shot.get("related_node_id", ""))
        row[2].text = str(shot.get("shot_type", ""))
        row[3].text = str(shot.get("camera_movement", ""))
        row[4].text = str(shot.get("duration_sec", ""))
        row[5].text = str(shot.get("visual_description", ""))
        row[6].text = str(shot.get("dialogue_or_sfx", ""))
        row[7].text = str(shot.get("prompt_draft", ""))
        row[8].text = str(shot.get("shooting_note", ""))
    doc.add_paragraph(f"\u9884\u4f30\u603b\u65f6\u957f: {storyboard.get('estimated_total_duration_sec', 0)} \u79d2")
    doc.add_paragraph(
        "\u5bfc\u51fa\u524d\u68c0\u67e5\u6e05\u5355: "
        + (", ".join(storyboard.get("export_ready_checklist", [])) if storyboard.get("export_ready_checklist") else "\u65e0")
    )

    doc.add_heading("\u89c6\u9891\u4efb\u52a1\u6458\u8981", level=2)
    for line in [
        f"\u5f53\u524d\u4efb\u52a1ID: {video_lab.get('task_id', '') or '-'}",
        f"\u5f53\u524d\u4efb\u52a1\u72b6\u6001: {video_lab.get('task_status', '') or '-'}",
        f"\u5f53\u524d\u89c6\u9891\u94fe\u63a5: {video_lab.get('video_url', '') or '-'}",
        f"\u957f\u89c6\u9891\u603b\u65f6\u957f: {video_lab.get('total_duration', 0)} \u79d2",
        f"\u6587\u4ef6\u540d\u524d\u7f00: {video_lab.get('filename_prefix', '') or '-'}",
    ]:
        doc.add_paragraph(line)
    if video_lab.get("long_segments"):
        for segment in video_lab.get("long_segments", []):
            doc.add_paragraph(
                f"\u7b2c{segment.get('index', '-')}段 | \u65f6\u957f: {segment.get('duration', 0)} \u79d2 | Task ID: {segment.get('task_id', '') or '-'} | \u72b6\u6001: {segment.get('task_status', '') or '-'}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("\u62c6\u6bb5\u4efb\u52a1: \u65e0", style="List Bullet")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _register_pdf_font() -> str:
    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        return "STSong-Light"
    except Exception:
        return "Helvetica"


def _draw_wrapped(
    c: canvas.Canvas, text: str, font_name: str, font_size: int, x: float, y: float, width: float
) -> float:
    lines = simpleSplit(str(text), font_name, font_size, width)
    for line in lines:
        c.drawString(x, y, line)
        y -= font_size + 4
    return y


def _build_pdf(payload: Dict[str, Any]) -> BytesIO:
    data = _normalize_export_payload(payload)
    project = data["project"]
    story_card = data["story_card"] or {}
    workshop = data["workshop"] or {}
    storyboard = data["storyboard"] or {}
    video_lab = data["video_lab"] or _default_project_state()["video_lab"]

    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    font_name = _register_pdf_font()

    def ensure_space(y_pos: float, need: float = 40.0) -> float:
        if y_pos < need:
            c.showPage()
            c.setFont(font_name, 11)
            return height - 50
        return y_pos

    def draw_section(title: str, items: List[str], y_pos: float) -> float:
        y_pos = ensure_space(y_pos, 60)
        c.setFont(font_name, 13)
        c.drawString(40, y_pos, title)
        y_pos -= 20
        c.setFont(font_name, 11)
        for item in items:
            y_pos = ensure_space(y_pos)
            y_pos = _draw_wrapped(c, item, font_name, 11, 50, y_pos, width - 90)
        return y_pos - 8

    y = height - 50
    c.setFont(font_name, 16)
    c.drawString(40, y, "\u77ed\u5267\u9879\u76ee\u5bfc\u51fa")
    y -= 30

    y = draw_section(
        "\u9879\u76ee\u4fe1\u606f",
        [
            f"\u9879\u76ee\u540d\u79f0: {project.get('name') or '未命名项目'}",
            f"\u521b\u5efa\u4eba: {project.get('creator', '') or '-'}",
            f"\u9879\u76ee\u63cf\u8ff0: {project.get('description', '') or '-'}",
            f"\u521b\u5efa\u65f6\u95f4: {project.get('created_at', '') or '-'}",
            f"\u66f4\u65b0\u65f6\u95f4: {project.get('updated_at', '') or '-'}",
            f"\u5bfc\u51fa\u65f6\u95f4: {data.get('exported_at', '') or '-'}",
            f"\u5f53\u524d\u6a21\u578b: {data.get('current_provider', '') or '-'}",
        ],
        y,
    )

    y = draw_section(
        "\u6545\u4e8b\u5361",
        [
            f"Logline: {story_card.get('logline', '') or '-'}",
            f"\u4e3b\u9898: {story_card.get('theme', '') or '-'}",
            f"\u57fa\u8c03: {story_card.get('tone', '') or '-'}",
            f"\u7ed3\u6784\u6a21\u677f: {story_card.get('structure_template', '') or '-'}",
            f"\u6838\u5fc3\u51b2\u7a81: {story_card.get('core_conflict', '') or '-'}",
            f"\u94a9\u5b50: {story_card.get('hook', '') or '-'}",
            f"\u7ed3\u5c40\u7c7b\u578b: {story_card.get('ending_type', '') or '-'}",
            "\u7ed3\u6784\u951a\u70b9: " + (", ".join(story_card.get("anchor_points", [])) if story_card.get("anchor_points") else "\u65e0"),
        ],
        y,
    )

    role_lines = [
        f"- {character.get('name') or '未命名角色'} | \u6807\u7b7e: {', '.join(character.get('tags', [])) or '-'} | \u52a8\u673a: {character.get('motivation', '') or '-'} | \u5f27\u5149: {character.get('arc', '') or '-'}"
        for character in workshop.get("characters", [])
    ] or ["- \u65e0"]
    y = draw_section("\u89d2\u8272\u8bbe\u5b9a", role_lines, y)

    relation_lines = [
        f"- {relationship.get('from', '-')} -> {relationship.get('to', '-')} | \u7c7b\u578b: {relationship.get('type', '') or '-'} | \u51b2\u7a81: {relationship.get('tension', '') or '-'}"
        for relationship in workshop.get("relationships", [])
    ] or ["- \u65e0"]
    y = draw_section("\u89d2\u8272\u5173\u7cfb", relation_lines, y)

    plot_lines = []
    for node in workshop.get("plot_nodes", []):
        plot_lines.extend(
            [
                f"- {node.get('id', '')} [{node.get('template_stage', '')}]",
                f"  \u6458\u8981: {node.get('summary', '') or '-'}",
                f"  \u5730\u70b9: {node.get('location', '') or '-'}",
                f"  \u52a8\u4f5c: {node.get('action_draft', '') or '-'}",
                f"  \u5bf9\u767d: {_dialogue_text(node.get('dialogue_draft'))}",
                f"  \u60c5\u611f\u53d8\u5316: {node.get('emotion_shift', '') or '-'}",
                f"  \u4e00\u81f4\u6027\u68c0\u67e5: {node.get('consistency_check', '') or '-'}",
            ]
        )
    if not plot_lines:
        plot_lines = ["- \u65e0"]
    y = draw_section("\u5267\u60c5\u8282\u70b9", plot_lines, y)

    timeline_lines = [f"{index}. {node_id}" for index, node_id in enumerate(workshop.get("timeline_view", []), start=1)] or ["- \u65e0"]
    y = draw_section("\u65f6\u95f4\u7ebf\u987a\u5e8f", timeline_lines, y)

    storyboard_lines = []
    for shot in storyboard.get("storyboards", []):
        storyboard_lines.extend(
            [
                f"{shot.get('shot_id', '')} | \u8282\u70b9: {shot.get('related_node_id', '')} | \u666f\u522b: {shot.get('shot_type', '')} | \u8fd0\u955c: {shot.get('camera_movement', '')} | \u65f6\u957f: {shot.get('duration_sec', 0)}\u79d2",
                f"\u753b\u9762: {shot.get('visual_description', '') or '-'}",
                f"\u5bf9\u767d/\u97f3\u6548: {shot.get('dialogue_or_sfx', '') or '-'}",
                f"\u63d0\u793a\u8bcd: {shot.get('prompt_draft', '') or '-'}",
                f"\u5907\u6ce8: {shot.get('shooting_note', '') or '-'}",
            ]
        )
    if not storyboard_lines:
        storyboard_lines = ["- \u65e0"]
    storyboard_lines.append(f"\u9884\u4f30\u603b\u65f6\u957f: {storyboard.get('estimated_total_duration_sec', 0)} \u79d2")
    storyboard_lines.append(
        "\u68c0\u67e5\u6e05\u5355: " + (", ".join(storyboard.get("export_ready_checklist", [])) if storyboard.get("export_ready_checklist") else "\u65e0")
    )
    y = draw_section("\u5206\u955c\u8868", storyboard_lines, y)

    video_lines = [
        f"\u5f53\u524d\u4efb\u52a1ID: {video_lab.get('task_id', '') or '-'}",
        f"\u5f53\u524d\u4efb\u52a1\u72b6\u6001: {video_lab.get('task_status', '') or '-'}",
        f"\u5f53\u524d\u89c6\u9891\u94fe\u63a5: {video_lab.get('video_url', '') or '-'}",
        f"\u957f\u89c6\u9891\u603b\u65f6\u957f: {video_lab.get('total_duration', 0)} \u79d2",
        f"\u6587\u4ef6\u540d\u524d\u7f00: {video_lab.get('filename_prefix', '') or '-'}",
    ]
    if video_lab.get("long_segments"):
        for segment in video_lab.get("long_segments", []):
            video_lines.append(
                f"\u7b2c{segment.get('index', '-')}段 | \u65f6\u957f: {segment.get('duration', 0)} \u79d2 | Task ID: {segment.get('task_id', '') or '-'} | \u72b6\u6001: {segment.get('task_status', '') or '-'}"
            )
    else:
        video_lines.append("\u62c6\u6bb5\u4efb\u52a1: \u65e0")
    y = draw_section("\u89c6\u9891\u4efb\u52a1\u6458\u8981", video_lines, y)

    c.save()
    buffer.seek(0)
    return buffer


@app.get("/")
def index() -> str:
    return render_template("index.html")


@app.get("/studio")
def studio() -> str:
    return render_template("studio.html")


@app.get("/visual")
def visual() -> str:
    return render_template("visual.html")


@app.get("/export-center")
def export_center() -> str:
    return render_template("export_center.html")


@app.get("/video-lab")
def video_lab() -> str:
    return render_template("video_lab.html")


@app.get("/api/providers")
def get_providers():
    providers_list = []
    for key, config in MODEL_PROVIDERS.items():
        if config.get("requires_qiniu_credential"):
            has_credential = _has_qiniu_text_credential()
        else:
            has_credential = bool(config.get("api_key"))

        providers_list.append({
            "id": key,
            "name": config["name"],
            "model": config["model"],
            "has_api_key": has_credential,
            "is_default": key == DEFAULT_PROVIDER
        })
    return jsonify({"ok": True, "providers": providers_list})


@app.get("/api/projects")
def get_projects():
    try:
        with _get_db_conn() as conn:
            projects = _list_projects(conn)
            return jsonify({"ok": True, "projects": projects})
    except Exception as e:  # noqa: BLE001
        return jsonify({"ok": False, "error": str(e)}), 500


@app.post("/api/projects")
def create_project_api():
    req_json = request.get_json(silent=True) or {}
    name = str(req_json.get("name", "")).strip()
    creator = str(req_json.get("creator", "")).strip()
    description = str(req_json.get("description", "")).strip()
    state = req_json.get("state")

    if not name:
        return jsonify({"ok": False, "error": "Project name is required."}), 400

    try:
        with _get_db_conn() as conn:
            project = _create_project(
                conn,
                name=name,
                creator=creator,
                description=description,
                state=state if isinstance(state, dict) else None,
                cover_image=str(req_json.get("cover_image", ""))[:200000],
                last_provider=str(req_json.get("last_provider", ""))[:64],
            )
            return jsonify({"ok": True, "project": project})
    except Exception as e:  # noqa: BLE001
        return jsonify({"ok": False, "error": str(e)}), 500


@app.get("/api/projects/<int:project_id>")
def get_project(project_id: int):
    try:
        with _get_db_conn() as conn:
            data = _get_project_with_state(conn, project_id)
            if not data:
                return jsonify({"ok": False, "error": "Project not found."}), 404
            return jsonify({"ok": True, **data})
    except Exception as e:  # noqa: BLE001
        return jsonify({"ok": False, "error": str(e)}), 500


@app.put("/api/projects/<int:project_id>")
def update_project(project_id: int):
    req_json = request.get_json(silent=True) or {}
    allowed_fields = {"name", "creator", "description", "cover_image", "last_provider"}
    try:
        with _get_db_conn() as conn:
            existing = conn.execute(
                "SELECT id FROM projects WHERE id = ? AND deleted = 0",
                (project_id,),
            ).fetchone()
            if not existing:
                return jsonify({"ok": False, "error": "Project not found."}), 404

            updates: List[str] = []
            params: List[Any] = []
            for field in allowed_fields:
                if field in req_json:
                    updates.append(f"{field} = ?")
                    if field in {"cover_image", "description"}:
                        params.append(str(req_json.get(field, ""))[:200000])
                    elif field == "last_provider":
                        params.append(str(req_json.get(field, ""))[:64])
                    else:
                        params.append(str(req_json.get(field, ""))[:255])

            now = _utc_now_iso()
            updates.append("updated_at = ?")
            params.append(now)
            params.append(project_id)
            conn.execute(
                f"UPDATE projects SET {', '.join(updates)} WHERE id = ?",
                tuple(params),
            )

            if "state" in req_json:
                state_obj = req_json.get("state")
                if not isinstance(state_obj, dict):
                    return jsonify({"ok": False, "error": "state must be an object."}), 400
                normalized_state = _normalize_project_state(state_obj)
                conn.execute(
                    """
                    INSERT INTO project_states(project_id, state_json, updated_at)
                    VALUES (?, ?, ?)
                    ON CONFLICT(project_id) DO UPDATE SET
                        state_json = excluded.state_json,
                        updated_at = excluded.updated_at
                    """,
                    (project_id, json.dumps(normalized_state, ensure_ascii=False), now),
                )

            conn.commit()
            data = _get_project_with_state(conn, project_id)
            return jsonify({"ok": True, **(data or {})})
    except Exception as e:  # noqa: BLE001
        return jsonify({"ok": False, "error": str(e)}), 500


@app.delete("/api/projects/<int:project_id>")
def delete_project(project_id: int):
    try:
        with _get_db_conn() as conn:
            existing = conn.execute(
                "SELECT id FROM projects WHERE id = ? AND deleted = 0",
                (project_id,),
            ).fetchone()
            if not existing:
                return jsonify({"ok": False, "error": "Project not found."}), 404

            now = _utc_now_iso()
            conn.execute(
                "UPDATE projects SET deleted = 1, updated_at = ? WHERE id = ?",
                (now, project_id),
            )
            conn.commit()

            fallback = _ensure_default_project(conn)
            projects = _list_projects(conn)
            return jsonify({
                "ok": True,
                "deleted_project_id": project_id,
                "fallback_project": fallback,
                "projects": projects,
            })
    except Exception as e:  # noqa: BLE001
        return jsonify({"ok": False, "error": str(e)}), 500


@app.get("/api/projects/<int:project_id>/state")
def get_project_state(project_id: int):
    try:
        with _get_db_conn() as conn:
            data = _get_project_with_state(conn, project_id)
            if not data:
                return jsonify({"ok": False, "error": "Project not found."}), 404
            return jsonify({"ok": True, "project_id": project_id, "state": data["state"]})
    except Exception as e:  # noqa: BLE001
        return jsonify({"ok": False, "error": str(e)}), 500


@app.put("/api/projects/<int:project_id>/state")
@app.post("/api/projects/<int:project_id>/state")
def put_project_state(project_id: int):
    req_json = request.get_json(silent=True) or {}
    state = req_json.get("state")
    if not isinstance(state, dict):
        return jsonify({"ok": False, "error": "state must be an object."}), 400
    normalized_state = _normalize_project_state(state)

    try:
        with _get_db_conn() as conn:
            existing = conn.execute(
                "SELECT id FROM projects WHERE id = ? AND deleted = 0",
                (project_id,),
            ).fetchone()
            if not existing:
                return jsonify({"ok": False, "error": "Project not found."}), 404

            now = _utc_now_iso()
            conn.execute(
                """
                INSERT INTO project_states(project_id, state_json, updated_at)
                VALUES (?, ?, ?)
                ON CONFLICT(project_id) DO UPDATE SET
                    state_json = excluded.state_json,
                    updated_at = excluded.updated_at
                """,
                (project_id, json.dumps(normalized_state, ensure_ascii=False), now),
            )
            conn.execute(
                "UPDATE projects SET updated_at = ? WHERE id = ?",
                (now, project_id),
            )
            if "cover_image" in req_json:
                conn.execute(
                    "UPDATE projects SET cover_image = ? WHERE id = ?",
                    (str(req_json.get("cover_image", ""))[:200000], project_id),
                )
            if "last_provider" in req_json:
                conn.execute(
                    "UPDATE projects SET last_provider = ? WHERE id = ?",
                    (str(req_json.get("last_provider", ""))[:64], project_id),
                )
            conn.commit()
            return jsonify({"ok": True, "project_id": project_id, "updated_at": now})
    except Exception as e:  # noqa: BLE001
        return jsonify({"ok": False, "error": str(e)}), 500


@app.post("/api/agent/compare")
def compare_providers():
    req_json = request.get_json(silent=True) or {}
    stage = req_json.get("stage")
    payload = req_json.get("payload", {})
    providers = req_json.get("providers", [])

    if not providers:
        return jsonify({"error": "No providers specified for comparison."}), 400

    results = {}
    errors = {}

    for provider in providers:
        try:
            if stage == "story_engine":
                result = _call_provider_json(
                    provider,
                    "你是专业短剧编剧策划，擅长结构化输出。",
                    _story_engine_prompt(payload),
                )
            elif stage == "workshop":
                result = _call_provider_json(
                    provider,
                    "你是专业短剧编剧，擅长角色与情节构建，并做一致性检查。",
                    _workshop_prompt(payload),
                )
            elif stage == "storyboard":
                result = _call_provider_json(
                    provider,
                    "你是分镜导演，擅长把剧情拆成可拍摄镜头。",
                    _storyboard_prompt(payload),
                )
            else:
                continue

            if stage == "story_engine":
                result = _normalize_story_engine_result(result)
            elif stage == "workshop":
                result = _normalize_workshop_result(result)
            elif stage == "storyboard":
                result = _normalize_storyboard_result(result)

            results[provider] = result
        except Exception as e:
            errors[provider] = str(e)

    return jsonify({
        "ok": True,
        "stage": stage,
        "results": results,
        "errors": errors
    })


@app.post("/api/agent/run")
def run_agent_stage():
    req_json = request.get_json(silent=True) or {}
    stage = req_json.get("stage")
    payload = req_json.get("payload", {})
    provider = req_json.get("provider", DEFAULT_PROVIDER)

    if stage not in {"story_engine", "workshop", "storyboard", "command", "export"}:
        return jsonify({"error": "Unsupported stage."}), 400

    try:
        if stage == "story_engine":
            result = _call_provider_json(
                provider,
                "你是专业短剧编剧策划，擅长结构化输出。",
                _story_engine_prompt(payload),
            )
        elif stage == "workshop":
            result = _call_provider_json(
                provider,
                "你是专业短剧编剧，擅长角色与情节构建，并做一致性检查。",
                _workshop_prompt(payload),
            )
        elif stage == "storyboard":
            result = _call_provider_json(
                provider,
                "你是分镜导演，擅长把剧情拆成可拍摄镜头。",
                _storyboard_prompt(payload),
            )
        elif stage == "command":
            result = _call_provider_json(
                provider,
                "你是编剧助手，负责执行自然语言编辑命令并保持一致性。",
                _command_prompt(payload),
            )
        else:
            result = _export_markdown(payload)

        if stage == "story_engine":
            result = _normalize_story_engine_result(result)
        elif stage == "workshop":
            result = _normalize_workshop_result(result)
        elif stage == "storyboard":
            result = _normalize_storyboard_result(result)
        elif stage == "command":
            result = _normalize_command_result(result)

        return jsonify({"ok": True, "stage": stage, "provider": provider, "result": result})
    except requests.HTTPError as e:
        detail: Optional[str] = None
        if e.response is not None:
            detail = e.response.text
        return jsonify({"ok": False, "error": "Model API request failed", "detail": detail}), 502
    except Exception as e:  # noqa: BLE001
        return jsonify({"ok": False, "error": str(e)}), 500


@app.post("/api/export/docx")
def export_docx():
    req_json = request.get_json(silent=True) or {}
    payload = req_json.get("payload", req_json)
    try:
        file_obj = _build_docx(payload)
        return send_file(
            file_obj,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="ai_short_drama_export.docx",
        )
    except Exception as e:  # noqa: BLE001
        return jsonify({"ok": False, "error": str(e)}), 500


@app.post("/api/export/pdf")
def export_pdf():
    req_json = request.get_json(silent=True) or {}
    payload = req_json.get("payload", req_json)
    try:
        file_obj = _build_pdf(payload)
        return send_file(
            file_obj,
            mimetype="application/pdf",
            as_attachment=True,
            download_name="ai_short_drama_export.pdf",
        )
    except Exception as e:  # noqa: BLE001
        return jsonify({"ok": False, "error": str(e)}), 500


@app.post("/api/video/script")
def generate_video_script():
    req_json = request.get_json(silent=True) or {}
    payload = req_json.get("payload", req_json)
    provider = req_json.get("provider", DEFAULT_PROVIDER)

    try:
        script = _call_provider_text(
            provider,
            "你是电影短剧导演和提示词工程师，擅长输出可直接用于视频生成的文本。",
            _video_script_prompt(payload),
        )
        return jsonify({"ok": True, "script": script, "provider": provider})
    except requests.HTTPError as e:
        detail: Optional[str] = None
        if e.response is not None:
            detail = e.response.text
        return jsonify({"ok": False, "error": "Model API request failed", "detail": detail}), 502
    except Exception as e:  # noqa: BLE001
        return jsonify({"ok": False, "error": str(e)}), 500


@app.post("/api/video/create-task")
def create_video_task():
    req_json = request.get_json(silent=True) or {}
    payload = req_json.get("payload", req_json)
    try:
        result = _create_video_task(payload)
        return jsonify({"ok": True, "result": result})
    except requests.HTTPError as e:
        detail: Optional[str] = None
        if e.response is not None:
            detail = e.response.text
        return jsonify({"ok": False, "error": "Video task creation failed", "detail": detail}), 502
    except Exception as e:  # noqa: BLE001
        return jsonify({"ok": False, "error": str(e)}), 500


@app.post("/api/video/create-long-task")
def create_long_video_task():
    """Create multiple short video tasks to approximate a long video.

    前端传入：
    - prompt: 整体视频提示词
    - model / size: 万相模型与分辨率
    - total_duration: 期望总时长（秒，可大于10）
    - segment_duration: 单段目标时长（可选，默认10，最大不会超过10）
    - provider: 用于续写提示词的对话模型提供商（可选，默认 DEFAULT_PROVIDER）
    """
    req_json = request.get_json(silent=True) or {}
    payload = req_json.get("payload", req_json)

    prompt = str(payload.get("prompt", "")).strip()
    if not prompt:
        return jsonify({"ok": False, "error": "Video prompt is required."}), 400

    try:
        total_duration = int(payload.get("total_duration", 0))
    except (TypeError, ValueError):
        total_duration = 0

    if total_duration <= 0:
        return jsonify({"ok": False, "error": "total_duration must be a positive integer."}), 400

    try:
        segment_duration = int(payload.get("segment_duration", 10))
    except (TypeError, ValueError):
        segment_duration = 10

    segment_duration = max(1, min(segment_duration, 10))

    provider = payload.get("provider") or DEFAULT_PROVIDER

    try:
        # 1) 先由对话模型规划每一段的文本提示词
        segments_plan = _extend_video_prompts(
            total_duration=total_duration,
            segment_duration=segment_duration,
            base_prompt=prompt,
            provider=provider,
        )

        # 2) 针对每一段调用万相创建视频任务
        video_model = payload.get("model", QINIU_VIDEO_MODEL)
        size = payload.get("size", "1280*720")
        prompt_extend = bool(payload.get("prompt_extend", True))

        segments_result: List[Dict[str, Any]] = []
        for seg in segments_plan:
            task_payload = {
                "prompt": seg["prompt"],
                "model": video_model,
                "size": size,
                "duration": int(seg["duration"]),
                "prompt_extend": prompt_extend,
            }
            task_raw = _create_video_task(task_payload)
            output = task_raw.get("output", {}) if isinstance(task_raw, dict) else {}
            segments_result.append(
                {
                    "index": seg["index"],
                    "duration": seg["duration"],
                    "prompt": seg["prompt"],
                    "task_id": output.get("task_id", ""),
                    "task_status": output.get("task_status", "PENDING"),
                }
            )

        return jsonify(
            {
                "ok": True,
                "result": {
                    "total_duration": total_duration,
                    "segment_duration": segment_duration,
                    "provider": provider,
                    "segments": segments_result,
                },
            }
        )
    except requests.HTTPError as e:
        detail: Optional[str] = None
        if e.response is not None:
            detail = e.response.text
        return (
            jsonify({"ok": False, "error": "Long video task creation failed", "detail": detail}),
            502,
        )
    except Exception as e:  # noqa: BLE001
        return jsonify({"ok": False, "error": str(e)}), 500


@app.get("/api/video/task/<task_id>")
def get_video_task(task_id: str):
    try:
        result = _query_video_task(task_id)
        return jsonify({"ok": True, "result": result})
    except requests.HTTPError as e:
        detail: Optional[str] = None
        if e.response is not None:
            detail = e.response.text
        return jsonify({"ok": False, "error": "Video task query failed", "detail": detail}), 502
    except Exception as e:  # noqa: BLE001
        return jsonify({"ok": False, "error": str(e)}), 500


@app.post("/api/search/web")
def web_search_api():
    req_json = request.get_json(silent=True) or {}
    query = str(req_json.get("query", "")).strip()
    if not query:
        return jsonify({"ok": False, "error": "query is required"}), 400

    max_results = int(req_json.get("max_results", 10))
    search_type = str(req_json.get("search_type", "web")).strip() or "web"

    try:
        result = _qiniu_web_search(query=query, max_results=max_results, search_type=search_type)
        return jsonify({"ok": True, "result": result})
    except requests.HTTPError as e:
        detail: Optional[str] = None
        if e.response is not None:
            detail = e.response.text
        return jsonify({"ok": False, "error": "Qiniu web search failed", "detail": detail}), 502
    except Exception as e:  # noqa: BLE001
        return jsonify({"ok": False, "error": str(e)}), 500


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8000, debug=True)
