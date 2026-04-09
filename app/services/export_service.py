from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import simpleSplit
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

from app.utils.helpers import (
    _as_text,
    _default_project_state,
    _first_present,
    _string_list,
    _utc_now_iso,
)
from app.utils.normalizers import (
    _normalize_story_card,
    _normalize_storyboard_result,
    _normalize_video_lab_state,
    _normalize_workshop_result,
)


def _dialogue_text(dialogue: Any) -> str:
    values = _string_list(dialogue)
    return " / ".join(values) if values else "-"


def _register_pdf_font() -> str:
    """
    注册 PDF 中文字体。
    优先顺序：
    1. 项目根目录 fonts/NotoSansSC-Regular.ttf
    2. Windows 常见中文字体
    找不到时直接报错，避免生成“看起来空白”的 PDF。
    """
    project_root = Path(__file__).resolve().parents[2]

    candidate_paths = [
        project_root / "fonts" / "NotoSansSC-Regular.ttf",
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
    ]

    font_name = "PDF_CJK_FONT"

    for font_path in candidate_paths:
        if font_path.exists():
            try:
                pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
                return font_name
            except Exception:
                continue

    raise RuntimeError(
        "未找到可用的中文 PDF 字体。"
        "请在项目根目录新建 fonts 文件夹，并放入 NotoSansSC-Regular.ttf，"
        "或确认系统存在可用中文字体。"
    )


def _draw_wrapped(
    pdf: canvas.Canvas,
    text: str,
    font_name: str,
    font_size: int,
    x: float,
    y: float,
    width: float,
) -> float:
    """
    在 PDF 中按宽度自动换行绘制文本，返回绘制后的新 y 坐标。
    """
    lines = simpleSplit(str(text), font_name, font_size, width)
    for line in lines:
        pdf.drawString(x, y, line)
        y -= font_size + 4
    return y


def _normalize_export_payload(payload: Dict[str, Any]) -> Dict[str, Any]:
    project = payload.get("project", {}) if isinstance(payload, dict) else {}
    if not isinstance(project, dict):
        project = {}

    return {
        "project": {
            "id": _as_text(project.get("id")),
            "name": _as_text(_first_present(project, "name", "project_name")) or "未命名项目",
            "creator": _as_text(project.get("creator")),
            "description": _as_text(project.get("description")),
            "created_at": _as_text(project.get("created_at")),
            "updated_at": _as_text(project.get("updated_at")),
        },
        "current_provider": _as_text(
            payload.get("current_provider") if isinstance(payload, dict) else ""
        )
        or _as_text(payload.get("provider") if isinstance(payload, dict) else "")
        or _as_text(project.get("last_provider")),
        "exported_at": _as_text(payload.get("exported_at") if isinstance(payload, dict) else "") or _utc_now_iso(),
        "story_card": _normalize_story_card(payload.get("story_card") if isinstance(payload, dict) else None),
        "workshop": _normalize_workshop_result(payload.get("workshop") if isinstance(payload, dict) else None),
        "storyboard": _normalize_storyboard_result(payload.get("storyboard") if isinstance(payload, dict) else None),
        "video_lab": _normalize_video_lab_state(payload.get("video_lab") if isinstance(payload, dict) else None),
    }


def _export_markdown(payload: Dict[str, Any]) -> Dict[str, Any]:
    data = _normalize_export_payload(payload)
    project = data["project"]
    story_card = data["story_card"] or {}
    workshop = data["workshop"] or {}
    storyboard = data["storyboard"] or {}
    video_lab = data["video_lab"] or _default_project_state()["video_lab"]

    lines: List[str] = []

    lines.append("# AI短剧项目导出")
    lines.append("")

    lines.append("## 0. 项目信息")
    lines.append(f"- 项目名称: {project.get('name', '未命名项目')}")
    lines.append(f"- 创建人: {project.get('creator', '') or '-'}")
    lines.append(f"- 项目描述: {project.get('description', '') or '-'}")
    lines.append(f"- 创建时间: {project.get('created_at', '') or '-'}")
    lines.append(f"- 更新时间: {project.get('updated_at', '') or '-'}")
    lines.append(f"- 导出时间: {data.get('exported_at', '') or '-'}")
    lines.append(f"- 当前模型: {data.get('current_provider', '') or '-'}")
    lines.append("")

    lines.append("## 1. 故事卡")
    lines.append(f"- Logline: {story_card.get('logline', '') or '-'}")
    lines.append(f"- 主题: {story_card.get('theme', '') or '-'}")
    lines.append(f"- 基调: {story_card.get('tone', '') or '-'}")
    lines.append(f"- 结构模板: {story_card.get('structure_template', '') or '-'}")
    lines.append(f"- 核心冲突: {story_card.get('core_conflict', '') or '-'}")
    lines.append(f"- 钩子: {story_card.get('hook', '') or '-'}")
    lines.append(f"- 结局类型: {story_card.get('ending_type', '') or '-'}")
    lines.append("- 结构锚点:")
    if story_card.get("anchor_points"):
        for index, point in enumerate(story_card.get("anchor_points", []), start=1):
            lines.append(f"  {index}. {point}")
    else:
        lines.append("  - 无")
    lines.append("")

    lines.append("## 2. 角色设定")
    if workshop.get("characters"):
        for character in workshop.get("characters", []):
            lines.append(
                f"- {character.get('name', '未命名角色')} | 标签: {', '.join(character.get('tags', [])) or '-'} | 动机: {character.get('motivation', '') or '-'} | 弧光: {character.get('arc', '') or '-'}"
            )
    else:
        lines.append("- 无")
    lines.append("")

    lines.append("## 3. 角色关系")
    if workshop.get("relationships"):
        for relationship in workshop.get("relationships", []):
            lines.append(
                f"- {relationship.get('from', '-')} -> {relationship.get('to', '-')} | 类型: {relationship.get('type', '') or '-'} | 冲突: {relationship.get('tension', '') or '-'}"
            )
    else:
        lines.append("- 无")
    lines.append("")

    lines.append("## 4. 剧情节点")
    if workshop.get("plot_nodes"):
        for node in workshop.get("plot_nodes", []):
            lines.append(f"- {node.get('id', '') or '-'} [{node.get('template_stage', '') or '-'}]")
            lines.append(f"  摘要: {node.get('summary', '') or '-'}")
            lines.append(f"  地点: {node.get('location', '') or '-'}")
            lines.append(f"  动作: {node.get('action_draft', '') or '-'}")
            lines.append(f"  对白: {_dialogue_text(node.get('dialogue_draft'))}")
            lines.append(f"  情绪变化: {node.get('emotion_shift', '') or '-'}")
            lines.append(f"  一致性检查: {node.get('consistency_check', '') or '-'}")
    else:
        lines.append("- 无")
    lines.append("")

    lines.append("## 5. 时间线顺序")
    if workshop.get("timeline_view"):
        for index, node_id in enumerate(workshop.get("timeline_view", []), start=1):
            lines.append(f"{index}. {node_id}")
    else:
        lines.append("- 无")
    lines.append("")

    lines.append("## 6. 分镜表")
    lines.append("| 镜头ID | 关联节点 | 景别 | 运镜 | 时长(秒) | 画面 | 对白/音效 | 提示词 | 备注 |")
    lines.append("|---|---|---|---|---|---|---|---|---|")
    if storyboard.get("storyboards"):
        for shot in storyboard.get("storyboards", []):
            lines.append(
                "| {shot_id} | {node} | {shot_type} | {move} | {duration} | {visual} | {sound} | {prompt} | {note} |".format(
                    shot_id=shot.get("shot_id", ""),
                    node=shot.get("related_node_id", ""),
                    shot_type=shot.get("shot_type", ""),
                    move=shot.get("camera_movement", ""),
                    duration=shot.get("duration_sec", ""),
                    visual=str(shot.get("visual_description", "")).replace("|", "\\|"),
                    sound=str(shot.get("dialogue_or_sfx", "")).replace("|", "\\|"),
                    prompt=str(shot.get("prompt_draft", "")).replace("|", "\\|"),
                    note=str(shot.get("shooting_note", "")).replace("|", "\\|"),
                )
            )
    lines.append("")
    lines.append(f"- 预估总时长: {storyboard.get('estimated_total_duration_sec', 0)} 秒")
    lines.append("- 导出前检查清单:")
    if storyboard.get("export_ready_checklist"):
        for item in storyboard.get("export_ready_checklist", []):
            lines.append(f"  - {item}")
    else:
        lines.append("  - 无")
    lines.append("")

    lines.append("## 7. 视频任务摘要")
    lines.append(f"- 当前任务ID: {video_lab.get('task_id', '') or '-'}")
    lines.append(f"- 当前任务状态: {video_lab.get('task_status', '') or '-'}")
    lines.append(f"- 当前视频链接: {video_lab.get('video_url', '') or '-'}")
    lines.append(f"- 长视频总时长: {video_lab.get('total_duration', 0)} 秒")
    lines.append(f"- 文件名前缀: {video_lab.get('filename_prefix', '') or '-'}")
    if video_lab.get("long_segments"):
        lines.append("- 拆段任务:")
        for segment in video_lab.get("long_segments", []):
            lines.append(
                f"  - 第{segment.get('index', '-')}段 | 时长: {segment.get('duration', 0)} 秒 | Task ID: {segment.get('task_id', '') or '-'} | 状态: {segment.get('task_status', '') or '-'}"
            )
    else:
        lines.append("- 拆段任务: 无")

    return {"markdown": "\n".join(lines)}


def _build_docx(payload: Dict[str, Any]) -> BytesIO:
    data = _normalize_export_payload(payload)
    project = data["project"]
    story_card = data["story_card"] or {}
    workshop = data["workshop"] or {}
    storyboard = data["storyboard"] or {}
    video_lab = data["video_lab"] or _default_project_state()["video_lab"]

    doc = Document()
    doc.add_heading("AI短剧项目导出", level=1)

    doc.add_heading("0. 项目信息", level=2)
    for line in [
        f"项目名称: {project.get('name', '未命名项目')}",
        f"创建人: {project.get('creator', '') or '-'}",
        f"项目描述: {project.get('description', '') or '-'}",
        f"创建时间: {project.get('created_at', '') or '-'}",
        f"更新时间: {project.get('updated_at', '') or '-'}",
        f"导出时间: {data.get('exported_at', '') or '-'}",
        f"当前模型: {data.get('current_provider', '') or '-'}",
    ]:
        doc.add_paragraph(line)

    doc.add_heading("1. 故事卡", level=2)
    for line in [
        f"Logline: {story_card.get('logline', '') or '-'}",
        f"主题: {story_card.get('theme', '') or '-'}",
        f"基调: {story_card.get('tone', '') or '-'}",
        f"结构模板: {story_card.get('structure_template', '') or '-'}",
        f"核心冲突: {story_card.get('core_conflict', '') or '-'}",
        f"钩子: {story_card.get('hook', '') or '-'}",
        f"结局类型: {story_card.get('ending_type', '') or '-'}",
    ]:
        doc.add_paragraph(line)

    doc.add_paragraph("结构锚点:")
    if story_card.get("anchor_points"):
        for point in story_card.get("anchor_points", []):
            doc.add_paragraph(point, style="List Bullet")
    else:
        doc.add_paragraph("无", style="List Bullet")

    doc.add_heading("2. 角色设定", level=2)
    if workshop.get("characters"):
        for character in workshop.get("characters", []):
            doc.add_paragraph(
                f"{character.get('name', '未命名角色')} | 标签: {', '.join(character.get('tags', [])) or '-'} | 动机: {character.get('motivation', '') or '-'} | 弧光: {character.get('arc', '') or '-'}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("无", style="List Bullet")

    doc.add_heading("3. 角色关系", level=2)
    if workshop.get("relationships"):
        for relationship in workshop.get("relationships", []):
            doc.add_paragraph(
                f"{relationship.get('from', '-')} -> {relationship.get('to', '-')} | 类型: {relationship.get('type', '') or '-'} | 冲突: {relationship.get('tension', '') or '-'}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("无", style="List Bullet")

    doc.add_heading("4. 剧情节点", level=2)
    if workshop.get("plot_nodes"):
        for node in workshop.get("plot_nodes", []):
            doc.add_paragraph(f"{node.get('id', '')} [{node.get('template_stage', '')}]", style="List Bullet")
            for detail in [
                f"摘要: {node.get('summary', '') or '-'}",
                f"地点: {node.get('location', '') or '-'}",
                f"动作: {node.get('action_draft', '') or '-'}",
                f"对白: {_dialogue_text(node.get('dialogue_draft'))}",
                f"情绪变化: {node.get('emotion_shift', '') or '-'}",
                f"一致性检查: {node.get('consistency_check', '') or '-'}",
            ]:
                doc.add_paragraph(detail)
    else:
        doc.add_paragraph("无", style="List Bullet")

    doc.add_heading("5. 时间线顺序", level=2)
    if workshop.get("timeline_view"):
        for node_id in workshop.get("timeline_view", []):
            doc.add_paragraph(node_id, style="List Number")
    else:
        doc.add_paragraph("无", style="List Bullet")

    doc.add_heading("6. 分镜表", level=2)
    table = doc.add_table(rows=1, cols=9)
    headers = ["镜头ID", "关联节点", "景别", "运镜", "时长", "画面", "对白/音效", "提示词", "备注"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    for shot in storyboard.get("storyboards", []):
        row = table.add_row().cells
        row[0].text = str(shot.get("shot_id", ""))
        row[1].text = str(shot.get("related_node_id", ""))
        row[2].text = str(shot.get("shot_type", ""))
        row[3].text = str(shot.get("camera_movement", ""))
        row[4].text = str(shot.get("duration_sec", ""))
        row[5].text = str(shot.get("visual_description", ""))
        row[6].text = str(shot.get("dialogue_or_sfx", ""))
        row[7].text = str(shot.get("prompt_draft", ""))
        row[8].text = str(shot.get("shooting_note", ""))

    doc.add_paragraph(f"预估总时长: {storyboard.get('estimated_total_duration_sec', 0)} 秒")
    doc.add_paragraph(
        "导出前检查清单: "
        + (", ".join(storyboard.get("export_ready_checklist", [])) if storyboard.get("export_ready_checklist") else "无")
    )

    doc.add_heading("7. 视频任务摘要", level=2)
    for line in [
        f"当前任务ID: {video_lab.get('task_id', '') or '-'}",
        f"当前任务状态: {video_lab.get('task_status', '') or '-'}",
        f"当前视频链接: {video_lab.get('video_url', '') or '-'}",
        f"长视频总时长: {video_lab.get('total_duration', 0)} 秒",
        f"文件名前缀: {video_lab.get('filename_prefix', '') or '-'}",
    ]:
        doc.add_paragraph(line)

    if video_lab.get("long_segments"):
        for segment in video_lab.get("long_segments", []):
            doc.add_paragraph(
                f"第{segment.get('index', '-')}段 | 时长: {segment.get('duration', 0)} 秒 | Task ID: {segment.get('task_id', '') or '-'} | 状态: {segment.get('task_status', '') or '-'}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("拆段任务: 无", style="List Bullet")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _build_pdf(payload: Dict[str, Any]) -> BytesIO:
    data = _normalize_export_payload(payload)
    project = data["project"]
    story_card = data["story_card"] or {}
    workshop = data["workshop"] or {}
    storyboard = data["storyboard"] or {}
    video_lab = data["video_lab"] or _default_project_state()["video_lab"]

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    font_name = _register_pdf_font()

    def ensure_space(y_pos: float, need: float = 40.0) -> float:
        if y_pos < need:
            pdf.showPage()
            pdf.setFont(font_name, 11)
            return height - 50
        return y_pos

    def draw_section(title: str, lines: List[str], y_pos: float) -> float:
        y_pos = ensure_space(y_pos, 60)
        pdf.setFont(font_name, 13)
        pdf.drawString(40, y_pos, title)
        y_pos -= 20

        pdf.setFont(font_name, 11)
        for line in lines:
            y_pos = ensure_space(y_pos)
            y_pos = _draw_wrapped(pdf, line, font_name, 11, 50, y_pos, width - 90)
        return y_pos - 8

    y = height - 50
    pdf.setFont(font_name, 16)
    pdf.drawString(40, y, "AI短剧项目导出")
    y -= 30

    y = draw_section(
        "0. 项目信息",
        [
            f"项目名称: {project.get('name', '未命名项目')}",
            f"创建人: {project.get('creator', '') or '-'}",
            f"项目描述: {project.get('description', '') or '-'}",
            f"创建时间: {project.get('created_at', '') or '-'}",
            f"更新时间: {project.get('updated_at', '') or '-'}",
            f"导出时间: {data.get('exported_at', '') or '-'}",
            f"当前模型: {data.get('current_provider', '') or '-'}",
        ],
        y,
    )

    y = draw_section(
        "1. 故事卡",
        [
            f"Logline: {story_card.get('logline', '') or '-'}",
            f"主题: {story_card.get('theme', '') or '-'}",
            f"基调: {story_card.get('tone', '') or '-'}",
            f"结构模板: {story_card.get('structure_template', '') or '-'}",
            f"核心冲突: {story_card.get('core_conflict', '') or '-'}",
            f"钩子: {story_card.get('hook', '') or '-'}",
            f"结局类型: {story_card.get('ending_type', '') or '-'}",
            "结构锚点: " + (", ".join(story_card.get("anchor_points", [])) if story_card.get("anchor_points") else "无"),
        ],
        y,
    )

    role_lines = [
        f"- {character.get('name', '未命名角色')} | 标签: {', '.join(character.get('tags', [])) or '-'} | 动机: {character.get('motivation', '') or '-'} | 弧光: {character.get('arc', '') or '-'}"
        for character in workshop.get("characters", [])
    ] or ["- 无"]
    y = draw_section("2. 角色设定", role_lines, y)

    relation_lines = [
        f"- {relationship.get('from', '-')} -> {relationship.get('to', '-')} | 类型: {relationship.get('type', '') or '-'} | 冲突: {relationship.get('tension', '') or '-'}"
        for relationship in workshop.get("relationships", [])
    ] or ["- 无"]
    y = draw_section("3. 角色关系", relation_lines, y)

    plot_lines: List[str] = []
    for node in workshop.get("plot_nodes", []):
        plot_lines.extend(
            [
                f"- {node.get('id', '')} [{node.get('template_stage', '')}]",
                f"  摘要: {node.get('summary', '') or '-'}",
                f"  地点: {node.get('location', '') or '-'}",
                f"  动作: {node.get('action_draft', '') or '-'}",
                f"  对白: {_dialogue_text(node.get('dialogue_draft'))}",
                f"  情绪变化: {node.get('emotion_shift', '') or '-'}",
                f"  一致性检查: {node.get('consistency_check', '') or '-'}",
            ]
        )
    if not plot_lines:
        plot_lines = ["- 无"]
    y = draw_section("4. 剧情节点", plot_lines, y)

    timeline_lines = [
        f"{index}. {node_id}"
        for index, node_id in enumerate(workshop.get("timeline_view", []), start=1)
    ] or ["- 无"]
    y = draw_section("5. 时间线顺序", timeline_lines, y)

    storyboard_lines: List[str] = []
    for shot in storyboard.get("storyboards", []):
        storyboard_lines.extend(
            [
                f"{shot.get('shot_id', '')} | 节点: {shot.get('related_node_id', '')} | 景别: {shot.get('shot_type', '')} | 运镜: {shot.get('camera_movement', '')} | 时长: {shot.get('duration_sec', 0)}秒",
                f"画面: {shot.get('visual_description', '') or '-'}",
                f"对白/音效: {shot.get('dialogue_or_sfx', '') or '-'}",
                f"提示词: {shot.get('prompt_draft', '') or '-'}",
                f"备注: {shot.get('shooting_note', '') or '-'}",
            ]
        )
    if not storyboard_lines:
        storyboard_lines = ["- 无"]
    storyboard_lines.append(f"预估总时长: {storyboard.get('estimated_total_duration_sec', 0)} 秒")
    storyboard_lines.append(
        "检查清单: "
        + (", ".join(storyboard.get("export_ready_checklist", [])) if storyboard.get("export_ready_checklist") else "无")
    )
    y = draw_section("6. 分镜表", storyboard_lines, y)

    video_lines = [
        f"当前任务ID: {video_lab.get('task_id', '') or '-'}",
        f"当前任务状态: {video_lab.get('task_status', '') or '-'}",
        f"当前视频链接: {video_lab.get('video_url', '') or '-'}",
        f"长视频总时长: {video_lab.get('total_duration', 0)} 秒",
        f"文件名前缀: {video_lab.get('filename_prefix', '') or '-'}",
    ]
    if video_lab.get("long_segments"):
        for segment in video_lab.get("long_segments", []):
            video_lines.append(
                f"第{segment.get('index', '-')}段 | 时长: {segment.get('duration', 0)} 秒 | Task ID: {segment.get('task_id', '') or '-'} | 状态: {segment.get('task_status', '') or '-'}"
            )
    else:
        video_lines.append("拆段任务: 无")
    y = draw_section("7. 视频任务摘要", video_lines, y)

    pdf.save()
    buffer.seek(0)
    return buffer
